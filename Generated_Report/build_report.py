#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build the complete B.Tech Final Year Project Report
"Development and Characterization of Bio-Composites from Waste Aquatic
 Biomass for Sustainable Insulation"  — NIT Srinagar, June 2026

Formatting per department guidelines:
  Times New Roman 12 pt, 1.5 line spacing, A4,
  margins L=38mm R=25mm T=25mm B=25mm,
  Roman numerals for preliminary pages, Arabic for main body,
  IEEE references (order of citation).
"""

import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHARTS = '/projects/sandbox/Final-Year-Project-Files/Generated_Report/charts'
PHOTOS = '/projects/sandbox/Final-Year-Project-Files/Generated_Report/photos'
OUTPUT = '/projects/sandbox/Final-Year-Project-Files/Generated_Report/FYP_Report_BioComposites.docx'

FONT = 'Times New Roman'
ACCENT = RGBColor(0x1F, 0x3B, 0x66)
GREEN  = RGBColor(0x2E, 0x6B, 0x2E)
GREY   = RGBColor(0x55, 0x55, 0x55)

# ─────────────────────────────────────────────────────────────────────────────
#  LOW-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), "true")
    trPr.append(th)

def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    r = run._r.find(qn('w:rPr'))
    return run

def set_number_format(section, fmt, start=1):
    """fmt: 'decimal' or 'lowerRoman'"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))

def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

def P(doc, text='', size=12, bold=False, italic=False, align=None,
      color=None, space_before=0, space_after=6, line=1.5, font=FONT,
      underline=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        if color is not None:
            r.font.color.rgb = color
    return p

def run_add(p, text, size=12, bold=False, italic=False, color=None, font=FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r

def heading(doc, text, level=1, space_before=14, space_after=8):
    """Uses built-in Heading styles (so the TOC field can pick them up),
    with Times New Roman font override."""
    sizes = {1: 15, 2: 13, 3: 12}
    style_name = f'Heading {min(level,3)}'
    p = doc.add_paragraph(style=style_name)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(sizes.get(level, 12))
    r.bold = True
    r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x20,0x20,0x20)
    return p

def add_toc_field(doc, switches=r'TOC \o "1-3" \h \z \u'):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = switches
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Right-click and select 'Update Field' to generate the Table of Contents."
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run._r.append(t); run._r.append(fld3)
    return p

def add_tof_field(doc, label='Figure'):
    """Table of Figures / Tables field based on caption label."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f'TOC \\h \\z \\c "{label}"'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = f"Right-click and 'Update Field' to generate the List of {label}s."
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run._r.append(t); run._r.append(fld3)
    return p

def seq_caption(doc, label, text, size=10.5, space_after=12):
    """Caption with a SEQ field so it is auto-numbered and picked up by the
    List of Figures/Tables. label = 'Figure' or 'Table'."""
    cp = doc.add_paragraph(style='Caption')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(space_after)
    cp.paragraph_format.line_spacing = 1.0
    r0 = cp.add_run(f'{label} ')
    r0.font.name = FONT; r0.font.size = Pt(size); r0.bold = True
    r0.italic = False; r0.font.color.rgb = RGBColor(0,0,0)
    # SEQ field
    run = cp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' SEQ {label} \\* ARABIC '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.name = FONT; run.font.size = Pt(size); run.bold = True
    rt = cp.add_run(f'.  {text}')
    rt.font.name = FONT; rt.font.size = Pt(size); rt.italic = False
    rt.font.color.rgb = RGBColor(0,0,0)
    return cp

def img(doc, path, width_in, cap_text, placeholder_label=None):
    """Insert a centered image (or placeholder) followed by an auto-numbered
    Figure caption."""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        label = placeholder_label or os.path.basename(path)
        # bordered placeholder
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        for side in ('top','bottom','left','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
            b.set(qn('w:space'),'8'); b.set(qn('w:color'),'AAAAAA')
            pbdr.append(b)
        pPr.append(pbdr)
        run_add(p, f'[ {label} \u2014 to be inserted ]',
                size=11, italic=True, color=GREY)
    seq_caption(doc, 'Figure', cap_text)

def tbl_cap(doc, text):
    """Auto-numbered Table caption placed above a table."""
    seq_caption(doc, 'Table', text, space_after=4)

def img_row(doc, items, total_width_in=6.2, cap_text=None):
    """Place multiple images side-by-side in a borderless table, with optional
    sub-labels, followed by a single auto-numbered Figure caption.
    items = list of (path, sublabel) tuples."""
    n = len(items)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    each_w = total_width_in / n
    for i, (path, sub) in enumerate(items):
        cell = t.rows[0].cells[i]
        cell.width = Inches(each_w)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        if os.path.exists(path):
            run = para.add_run()
            run.add_picture(path, width=Inches(each_w - 0.18))
        else:
            run_add(para, f'[ {sub or os.path.basename(path)} ]',
                    size=10, italic=True, color=GREY)
        if sub:
            sp = cell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_after = Pt(2)
            run_add(sp, sub, size=9.5, italic=True, color=GREY)
    # remove borders
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'none'); borders.append(e)
    tblPr.append(borders)
    if cap_text:
        seq_caption(doc, 'Figure', cap_text)

def make_table(doc, headers, rows, col_widths=None, header_bg='1F3B66',
               font_size=10.5, header_color='FFFFFF', zebra=True,
               highlight=None):
    """highlight: dict {(row_idx, col_idx): 'hexcolor'} for special cells"""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    # header
    hdr = t.rows[0].cells
    set_repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], header_bg)
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        r = para.add_run(h)
        r.font.name = FONT; r.font.size = Pt(font_size); r.bold = True
        r.font.color.rgb = RGBColor.from_string(header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # body
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            para = cells[ci].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            r = para.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(font_size)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra and ri % 2 == 1:
                set_cell_bg(cells[ci], 'F2F5FA')
            if highlight and (ri, ci) in highlight:
                set_cell_bg(cells[ci], highlight[(ri, ci)])
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return t

def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1F3B66')
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(4)

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
style_base(doc)

sec = doc.sections[0]
sec.page_height = Mm(297); sec.page_width = Mm(210)   # A4
sec.left_margin   = Mm(38)
sec.right_margin  = Mm(25)
sec.top_margin    = Mm(25)
sec.bottom_margin = Mm(25)

print("Document base configured.")
print("Builder part 1 (setup + helpers) loaded OK.")



# ═════════════════════════════════════════════════════════════════════════════
#  TITLE PAGES  (Outer + Inner, per Annexure-I format)
# ═════════════════════════════════════════════════════════════════════════════
TITLE_FULL = ("DEVELOPMENT AND CHARACTERIZATION OF BIO-COMPOSITES FROM "
              "WASTE AQUATIC BIOMASS FOR SUSTAINABLE INSULATION")

def title_page(doc):
    P(doc, TITLE_FULL, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
      color=ACCENT, space_before=18, space_after=18)
    P(doc, 'A PROJECT', size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_after=14)
    P(doc, 'Submitted in partial fulfillment of the requirements for',
      size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    P(doc, 'the award of the degree of', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    P(doc, 'BACHELOR OF TECHNOLOGY', size=14, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    P(doc, 'By', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    for nm, en in [("SUCHETAN KARLOOPIA", "2022BCHE019"),
                   ("MIRAN HAIDER", "2022BCHE027"),
                   ("AKSHITA SEN", "2022BCHE037")]:
        P(doc, nm, size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=0)
        P(doc, f'(Registration Number: {en})', size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    P(doc, 'Under the guidance', size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_before=8, space_after=0)
    P(doc, 'Of  Dr. Fasil Qayoom Mir', size=12, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    P(doc, 'DEPARTMENT OF CHEMICAL ENGINEERING', size=13, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=2)
    P(doc, 'NATIONAL INSTITUTE OF TECHNOLOGY', size=13, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=2)
    P(doc, 'SRINAGAR \u2013 190006 (INDIA)', size=12, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    P(doc, 'June, 2026', size=12, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# Outer title page
title_page(doc)
page_break(doc)
# Inner title page
title_page(doc)
# Copyright (on reverse side of inner title page)
P(doc, 'COPYRIGHT \u00a9 NIT SRINAGAR (J&K), INDIA, 2026', size=12, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=2)
P(doc, 'All rights reserved.', size=11, italic=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  CANDIDATE'S DECLARATION  (with Supervisor Certificate + Viva-Voce)
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "NATIONAL INSTITUTE OF TECHNOLOGY SRINAGAR (J&K)", size=12.5, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
P(doc, "CANDIDATES' DECLARATION", size=14, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
hrule(doc)
P(doc, "We hereby certify that the work which is being presented in the project "
  "titled \u201c" + TITLE_FULL.title() + "\u201d in partial fulfillment of the "
  "requirements for the award of the Degree of Bachelor of Technology and "
  "submitted in the Department of Chemical Engineering, National Institute of "
  "Technology Srinagar, is an authentic record of our own work carried out "
  "during a period from August 2025 to June 2026 under the supervision of "
  "Dr. Fasil Qayoom Mir, Head & Associate Professor, Department of Chemical "
  "Engineering, National Institute of Technology Srinagar.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "The matter presented in this project report has not been submitted by us "
  "for the award of any other degree of this or any other Institute/University.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=16)
for name, enrol in [("Suchetan Karloopia", "2022BCHE019"),
                    ("Miran Haider", "2022BCHE027"),
                    ("Akshita Sen", "2022BCHE037")]:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pp.paragraph_format.space_after = Pt(2)
    run_add(pp, "Sd/-", size=12)
    pp2 = doc.add_paragraph(); pp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pp2.paragraph_format.space_after = Pt(12)
    run_add(pp2, f"({name.upper()}) \u2014 {enrol}", size=12, bold=True)
P(doc, "This is to certify that the above statement made by the candidates is "
  "correct to the best of my knowledge.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=8, space_after=16)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pp.paragraph_format.space_after = Pt(2)
run_add(pp, "Sd/-", size=12)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pp.paragraph_format.space_after = Pt(0)
run_add(pp, "(Dr. Fasil Qayoom Mir)", size=12, bold=True)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pp.paragraph_format.space_after = Pt(16)
run_add(pp, "Head & Associate Professor", size=11)
P(doc, "The project Viva-Voce Examination of Suchetan Karloopia, Miran Haider "
  "and Akshita Sen has been held on ____________________.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=8, space_after=30)
pp = doc.add_paragraph()
tabs = pp.paragraph_format.tab_stops
tabs.add_tab_stop(Inches(4.3), WD_TAB_ALIGNMENT.LEFT)
run_add(pp, "Signature of Supervisor(s)\tSignature of External Examiner", size=11)
P(doc, "Date:  ______________", size=11, space_before=12, space_after=2)
page_break(doc)

print("Title pages + Copyright + Declaration/Certificate added.")



# ═════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "ACKNOWLEDGEMENT", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=12)
hrule(doc)
P(doc, "We wish to express our deepest gratitude and sincere appreciation to our "
  "supervisor, Dr. Fasil Qayoom Mir, Head and Associate Professor, Department of "
  "Chemical Engineering, NIT Srinagar, for his invaluable guidance, constant "
  "encouragement, and constructive criticism throughout the course of this "
  "project. His insight and expertise were instrumental in shaping this work, and "
  "it has been a privilege to work under his supervision.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "We are grateful to Prof. Tabassum Ara, Department of Chemistry, NIT "
  "Srinagar, for kindly providing access to the mixer-grinder facility used for "
  "biomass size reduction, and to Dr. Khalid Majid for facilitating access to the "
  "KD2 Pro Thermal Properties Analyzer for thermal conductivity measurements.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "We extend our thanks to the faculty and technical staff of the Department "
  "of Chemical Engineering for providing the laboratory facilities and assistance "
  "required to carry out the experimental work. We also acknowledge the support "
  "of the Department of Civil Engineering for access to the Baker Type K12 "
  "unconfined compression testing apparatus.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "Finally, we thank our families and friends for their unwavering support "
  "and motivation throughout our academic journey.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24)

pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_add(pp, "Suchetan Karloopia", size=12, bold=True)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_add(pp, "Miran Haider", size=12, bold=True)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_add(pp, "Akshita Sen", size=12, bold=True)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "ABSTRACT", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=12)
hrule(doc)
P(doc, "Buildings account for a large fraction of global energy use, and most of "
  "the insulation used to cut that demand is far from green. Expanded polystyrene "
  "(EPS), polyurethane (PU) foam and glass wool are either made from petroleum or "
  "are energy-intensive to manufacture, which has pushed researchers to look for "
  "sustainable materials made from renewable or waste resources. At the same time, "
  "the invasive weeds water hyacinth (Eichhornia crassipes) and water lily "
  "(Nymphaea spp.) continue to choke Dal Lake in Srinagar, a Ramsar wetland of "
  "roughly 18 km\u00b2 that is badly affected by eutrophication. These weeds are "
  "cleared in large quantities every year but are usually thrown away, so they "
  "form a steady and essentially free stream of waste biomass. This project links "
  "the two problems together by turning that waste biomass into bio-composite "
  "insulation panels.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "Four formulations were prepared from oven-dried biomass that had been "
  "ground and classified using ASTM E11 sieves: a coarse binder-less panel (S1), "
  "a fine binder-less panel (S2), and two fine panels held together with "
  "gelatinised food-grade corn starch at biomass:starch ratios of 90:10 (S3) and "
  "70:30 (S4). We deliberately avoided synthetic polymers and chemical "
  "crosslinkers altogether. Each panel was then tested for moisture content "
  "(ASTM D4442), bulk density (ASTM D1037), water absorption (ASTM D570), "
  "unconfined compressive strength (Baker Type K12 UCT apparatus) and thermal "
  "conductivity (KD2 Pro TR-3 transient line-source probe, ASTM D5334), with "
  "three replicates run for every test.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "The fine binder-less panel (S2) and the 90:10 starch panel (S3) reached "
  "insulation-grade thermal conductivities of 0.0577 and 0.0608 W/m\u00b7K "
  "respectively, both comfortably under the 0.065 W/m\u00b7K limit set by "
  "ASTM C168. The 90:10 panel (S3) also gave the highest unconfined compressive "
  "strength at 186.0 kPa, a 97 % gain over the binder-less fine panel, which "
  "points to 10 wt% corn starch as the sweet spot for binder content. The "
  "high-starch panel (S4, 70:30) was a more puzzling case: it absorbed the least "
  "water (280 %), yet it was also the weakest (26.8 kPa) and conducted heat about "
  "three times faster (0.1846 W/m\u00b7K). We attribute this to an excess of "
  "plasticised starch and a high residual moisture content of 41.18 %. Taken "
  "together, the results make the 90:10 biomass:starch panel (S3) the best "
  "all-round formulation, since it combines insulation-grade thermal performance "
  "with the highest strength, and they show that an invasive aquatic weed can "
  "realistically be turned into a workable, biodegradable building-insulation "
  "material.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
pk = doc.add_paragraph(); pk.paragraph_format.space_after = Pt(6)
run_add(pk, "Keywords: ", size=12, bold=True)
run_add(pk, "Water hyacinth; Water lily; Bio-composite; Corn starch binder; "
        "Thermal insulation; Unconfined compressive strength; Sustainable "
        "building materials; Waste valorisation.", size=12, italic=True)
page_break(doc)

print("Acknowledgement + Abstract added.")



# ═════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (auto field)
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "TABLE OF CONTENTS", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
add_toc_field(doc)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES  (auto field)
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "LIST OF FIGURES", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
add_tof_field(doc, 'Figure')
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES  (auto field)
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "LIST OF TABLES", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
add_tof_field(doc, 'Table')
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  LIST OF ABBREVIATIONS
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "LIST OF ABBREVIATIONS", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
abbr = [
    ("ASTM", "American Society for Testing and Materials"),
    ("WH", "Water Hyacinth (Eichhornia crassipes)"),
    ("WL", "Water Lily (Nymphaea spp.)"),
    ("WHF", "Water Hyacinth Fibre"),
    ("MC", "Moisture Content"),
    ("WA", "Water Absorption"),
    ("UCT", "Unconfined Compression Test"),
    ("SEM", "Scanning Electron Microscopy"),
    ("EPS", "Expanded Polystyrene"),
    ("XPS", "Extruded Polystyrene"),
    ("PU", "Polyurethane"),
    ("IFR", "Intumescent Flame Retardant"),
    ("PBS", "Poly(butylene succinate)"),
    ("LOI", "Limiting Oxygen Index"),
    ("RH", "Relative Humidity"),
    ("LCA", "Life Cycle Assessment"),
    ("IEEE", "Institute of Electrical and Electronics Engineers"),
    ("SD", "Standard Deviation"),
]
make_table(doc, ["Abbreviation", "Description"], abbr,
           col_widths=[1.8, 4.5], font_size=11, zebra=True)
page_break(doc)

# ═════════════════════════════════════════════════════════════════════════════
#  LIST OF SYMBOLS
# ═════════════════════════════════════════════════════════════════════════════
P(doc, "LIST OF SYMBOLS", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
symbols = [
    ("K", "Thermal conductivity", "W/m\u00b7K"),
    ("R", "Thermal resistivity", "\u00b0C\u00b7cm/W"),
    ("\u03c1", "Bulk density", "g/cm\u00b3"),
    ("\u03c3", "Compressive (deviator) stress", "kPa"),
    ("\u03b5", "Axial strain", "\u2013  or  %"),
    ("\u03b5f", "Failure strain", "%"),
    ("q\u1d64", "Unconfined compressive strength", "kPa"),
    ("S\u1d64", "Undrained shear strength (= q\u1d64/2)", "kPa"),
    ("A\u2080", "Initial cross-sectional area", "mm\u00b2"),
    ("A", "Corrected cross-sectional area", "mm\u00b2"),
    ("H\u2080", "Initial specimen height", "mm"),
    ("\u0394L", "Axial deformation", "mm"),
    ("F", "Applied compressive force", "N"),
    ("C", "Proving-ring calibration constant", "N/div"),
    ("W\u1d62", "Initial (wet) mass", "g"),
    ("W\u0066", "Final (oven-dry) mass", "g"),
    ("Syx", "Standard error of estimate (goodness-of-fit)", "\u2013"),
    ("n", "Number of replicates", "\u2013"),
]
make_table(doc, ["Symbol", "Description", "Unit"], symbols,
           col_widths=[1.2, 3.8, 1.3], font_size=11, zebra=True)
page_break(doc)

print("TOC + Lists + Abbreviations + Symbols added.")



# ═════════════════════════════════════════════════════════════════════════════
#  SECTION BREAK -> MAIN BODY (Arabic numerals from 1)
# ═════════════════════════════════════════════════════════════════════════════
main_sec = doc.add_section(WD_SECTION.NEW_PAGE)
main_sec.page_height = Mm(297); main_sec.page_width = Mm(210)
main_sec.left_margin = Mm(38); main_sec.right_margin = Mm(25)
main_sec.top_margin = Mm(25); main_sec.bottom_margin = Mm(25)
main_sec.footer.is_linked_to_previous = False

# ─────────────────────────────────────────────────────────────────────────────
#  CHAPTER 1 — INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 1", level=1, space_before=0)
heading(doc, "INTRODUCTION", level=1, space_before=2)

heading(doc, "1.1  Background", level=2)
P(doc, "Thermal insulation is one of the cheapest and most effective ways to cut "
  "energy use in buildings, which are responsible for a large slice of global "
  "energy consumption and greenhouse-gas emissions. The trouble is that most "
  "insulation on the market today \u2014 expanded polystyrene (EPS), extruded "
  "polystyrene (XPS), polyurethane (PU) foam, glass wool and mineral wool \u2014 "
  "is either petroleum-derived or energy-intensive to make. These products do not "
  "biodegrade, they pile up in landfills once their service life ends, and the "
  "organic foams among them burn readily and give off toxic gases when they do "
  "[1], [7]. With circular-economy thinking and carbon-footprint reduction now "
  "high on the agenda, attention has shifted towards insulation made from "
  "renewable and waste lignocellulosic resources.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Natural-fibre bio-composites are an attractive alternative for several "
  "reasons. They are renewable and biodegradable, they are light, their porous "
  "cellular structure gives them a naturally low thermal conductivity, and they "
  "can often be made cheaply from agricultural or aquatic waste. Researchers have "
  "tried a long list of lignocellulosic feedstocks for insulation \u2014 rice "
  "straw, corn cob, coir, hemp, flax, kenaf, sugarcane bagasse, pineapple leaf "
  "and date-palm waste, to name a few [4], [6]. Of the aquatic options, water "
  "hyacinth (Eichhornia crassipes) stands out because it is abundant, grows "
  "extremely fast and has a high cellulose content.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "How well a material insulates is captured mainly by its thermal "
  "conductivity, K, measured in watts per metre-kelvin (W/m\u00b7K). The lower the "
  "K, the better the insulator, because less heat passes through the material for "
  "a given temperature difference. Lightweight insulators get their low "
  "conductivity from the large amount of still air trapped inside a cellular or "
  "fibrous solid; air on its own conducts heat very poorly, at roughly "
  "0.026 W/m\u00b7K. Anything that pushes this trapped air out \u2014 "
  "densification, void filling, or above all moisture getting in \u2014 raises the "
  "effective conductivity and spoils the insulation. Water conducts heat at about "
  "0.6 W/m\u00b7K, some 23 times better than air, which is exactly why keeping "
  "moisture out keeps coming up as a problem in bio-based insulation.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "A bio-composite combines a natural fibre or particle reinforcement with a "
  "matrix, or binder, so that the resulting material performs better than either "
  "ingredient would on its own. Here the reinforcement is lignocellulosic aquatic "
  "biomass and the matrix is a biodegradable polysaccharide binder. The appeal of "
  "such composites lies in their renewability, biodegradability, low density and "
  "low embodied energy and carbon, along with the fact that they put a waste "
  "stream to good use. Their weak points are equally well known: they take up a "
  "lot of moisture, their mechanical strength is modest, and they are flammable. "
  "These are precisely the shortcomings that tuning particle size, binder type "
  "and binder content is meant to address, and they are what this project sets "
  "out to study in a systematic way.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "1.2  The Dal Lake Problem: An Invasive Weed as a Resource", level=2)
P(doc, "Dal Lake in Srinagar, Jammu & Kashmir, is one of the Himalayas' best-known "
  "urban lakes and a designated Ramsar wetland of roughly 18 km\u00b2. Over the "
  "last few decades it has declined badly, mainly through eutrophication fed by "
  "untreated sewage, the spread of the city into its catchment and heavy nutrient "
  "loading. One of the most obvious results of all that nutrient enrichment is the "
  "runaway growth of free-floating invasive weeds, chiefly water hyacinth "
  "(Eichhornia crassipes) and water lily (Nymphaea spp.). They spread into thick "
  "surface mats that starve the water of oxygen, cut out sunlight, wreck native "
  "fish breeding grounds, get in the way of boats and gradually turn open water "
  "into marsh.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Every year, mechanical de-weeding pulls large amounts of this biomass out "
  "of the lake, but the harvested material is treated as rubbish and either dumped "
  "or left to rot, releasing methane and feeding the same nutrients back into the "
  "water. The upshot is a continuous, free and otherwise troublesome supply of "
  "biomass. Turning this weed into something useful tackles two problems at once: "
  "it gives weed removal an economic payoff and an end use, and it provides a "
  "renewable raw material for sustainable insulation. That dual benefit is the "
  "central idea behind this project.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img(doc, f'{PHOTOS}/__missing_dallake.jpg', 5.6,
    "Dal Lake, Srinagar (\u2248 N 34.08\u00b0, E 74.85\u00b0) showing aquatic "
    "weed infestation \u2014 collection site of the raw biomass.",
    placeholder_label="Figure: Dal Lake collection site photograph")

heading(doc, "1.3  Motivation", level=2)
P(doc, "Three things together motivated this work. The first is the environmental "
  "and health baggage that comes with petroleum-based synthetic insulation. The "
  "second is the pressing need to manage and make use of the invasive weed that "
  "clogs Dal Lake. The third is recent evidence that water-hyacinth fibre can "
  "perform well as a low-thermal-conductivity natural material [2], [3]. By "
  "choosing a food-grade, biodegradable corn-starch binder instead of synthetic "
  "resins or chemical crosslinkers, we aimed to make a fully bio-based, low-cost "
  "and environmentally harmless insulation panel that can be produced with "
  "simple, scalable equipment.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "1.4  Problem Statement", level=2)
P(doc, "Waste aquatic biomass is a promising green substitute for synthetic "
  "insulation, but on its own it cannot be used directly. It is too weak "
  "mechanically, it soaks up too much moisture, and it catches fire easily. "
  "Overcoming these limitations calls for an engineered bio-composite \u2014 the "
  "right binder paired with the right processing route \u2014 that delivers enough "
  "structural integrity, dimensional stability and insulating performance while "
  "still staying biodegradable.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "1.5  Objectives", level=2)
P(doc, "The primary objectives of this project are:", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
  space_after=2)
for i, obj in enumerate([
    "To develop eco-friendly bio-composite insulation panels from waste aquatic "
    "biomass (water hyacinth and water lily) collected from Dal Lake.",
    "To fabricate composite variants differing in particle size and "
    "biomass-to-binder ratio using a natural, gelatinised corn-starch binder.",
    "To characterize the physical properties of the panels \u2014 moisture "
    "content, bulk density and water absorption \u2014 following ASTM standards.",
    "To evaluate the mechanical performance of the panels through unconfined "
    "compression testing and to determine the optimal binder content.",
    "To measure the thermal conductivity of the panels and assess their "
    "suitability as insulation-grade materials against the ASTM C168 threshold.",
    "To identify the best-balanced formulation by comparing mechanical, thermal "
    "and moisture-resistance performance, and to benchmark it against "
    "conventional insulation materials reported in the literature.",
]):
    pp = doc.add_paragraph(style='List Number')
    pp.paragraph_format.line_spacing = 1.5
    pp.paragraph_format.space_after = Pt(3)
    r = pp.add_run(obj); r.font.name = FONT; r.font.size = Pt(12)

heading(doc, "1.6  Scope of the Work", level=2)
P(doc, "This study is concerned with fabricating four bio-composite formulations "
  "at laboratory scale and characterising their physical, mechanical and thermal "
  "behaviour. The testing programme covers moisture content (ASTM D4442), bulk "
  "density (ASTM D1037), water absorption (ASTM D570), unconfined compressive "
  "strength (Baker Type K12 apparatus), thermal conductivity (KD2 Pro TR-3 probe, "
  "ASTM D5334) and scanning electron microscopy (SEM) of the binder-less (S1) and "
  "90:10 starch-bound (S3) panels to examine their microstructure. Flammability "
  "testing, hydrophobic surface treatments and a full life-cycle assessment are "
  "left for future work and are not covered here.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "1.7  Organisation of the Report", level=2)
P(doc, "The rest of the report is organised as follows. Chapter 1 has set out the "
  "background, motivation, problem statement and objectives. Chapter 2 reviews the "
  "literature on natural-fibre and water-hyacinth-based insulation and pins down "
  "the research gap. Chapter 3 covers the design, experimental setup and "
  "methodology, including the raw materials, binder preparation, sample "
  "fabrication and the test methods and equipment. Chapter 4 reports the results "
  "and discusses them alongside published work. Chapter 5 draws the findings "
  "together, states the main conclusions and outlines where the work could go "
  "next.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
page_break(doc)

print("Chapter 1 added.")



# ─────────────────────────────────────────────────────────────────────────────
#  CHAPTER 2 — LITERATURE REVIEW
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 2", level=1, space_before=0)
heading(doc, "LITERATURE REVIEW", level=1, space_before=2)

heading(doc, "2.1  Overview", level=2)
P(doc, "Plenty of research has looked at natural fibres and agricultural or "
  "aquatic wastes as raw materials for sustainable insulation composites. This "
  "chapter walks through the studies most relevant to the present work, paying "
  "particular attention to water-hyacinth-based composites, natural binder "
  "systems, fire-retardancy and what happens to these materials at end of life. "
  "It sets the performance benchmarks that this project is measured against and "
  "marks out the gap that the work aims to fill.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "2.2  Water-Hyacinth-Based Insulation Composites", level=2)
P(doc, "Salas-Ruiz et al. [2] studied binder-less particle boards made from "
  "water-hyacinth petioles and measured thermal conductivities as low as "
  "0.047 W/m\u00b7K for staple-fibre boards and around 0.065 W/m\u00b7K for pulp "
  "boards. Water absorption, though, was very high at 450\u2013555 %, which neatly "
  "captures the two sides of these boards: they insulate beautifully but are very "
  "sensitive to moisture. Philip and Rakendu [3] went a different route with a "
  "water-hyacinth\u2013cement composite, reaching a thermal conductivity of "
  "0.0765 W/m\u00b7K at a density of 0.47 g/cm\u00b3, with a flexural strength of "
  "about 0.35 MPa and water absorption close to 98 %. Jaktorn and Jiajitsawat [1] "
  "bonded water-hyacinth fibre with natural rubber latex and obtained very low "
  "conductivities of 0.0246\u20130.0305 W/m\u00b7K, but their boards still failed "
  "the relevant industrial standard on water absorption \u2014 once again, "
  "moisture resistance was the sticking point.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Anjani et al. [6] made water-hyacinth\u2013sugarcane-bagasse fibre\u2013"
  "epoxy composites for cool-box insulation, with their best configuration giving "
  "a thermal conductivity of about 0.1987 W/m\u00b7K and a maximum bending "
  "strength of 263 kgf/cm\u00b2. Chaireh et al. [5] worked on starch\u2013"
  "water-hyacinth foams for food packaging and found a 5 wt% water-hyacinth "
  "loading to be optimal, with the elastic modulus reaching about 232 MPa. A "
  "beeswax coating cut water absorption sharply in their work, which is doubly "
  "useful here: it shows starch is a workable natural binder for water hyacinth "
  "and that beeswax makes an effective hydrophobic treatment.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "2.3  Binder Systems, Surface Treatment and Fire Retardancy", level=2)
P(doc, "The binder has a big say in how a natural-fibre composite behaves, both "
  "mechanically and in the wet. Syamsuri et al. [8] showed that treating "
  "water-hyacinth fibre with alkali (NaOH) inside a cassava-starch bioplastic "
  "roughly quadrupled its tensile strength compared with untreated fibre, simply "
  "by improving how well the fibre stuck to the matrix. Abral et al. [10] sounded "
  "a note of caution, though: water-hyacinth\u2013polyester composites processed "
  "while wet came out mechanically weaker, and aggressive alkali treatment opened "
  "up micro-voids that hurt performance. The lesson is that controlling moisture "
  "during processing and curing really matters. On the fire-safety side, "
  "Suwanniroj and Suppakarn [9] used water-hyacinth fibre as a bio-based carbon "
  "source in an intumescent flame-retardant poly(butylene succinate) system and "
  "achieved a limiting oxygen index of 28.8 %, a UL-94 V-0 rating and a 53 % drop "
  "in peak heat-release rate.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "2.4  Conventional Insulation Benchmarks", level=2)
P(doc, "It helps to put bio-based numbers next to conventional insulation. Jeon "
  "et al. [7] measured a thermal conductivity of about 0.034 W/m\u00b7K for glass "
  "wool and noted that moisture getting in can push the conductivity of fibrous "
  "inorganic insulation up by roughly four times. Bio-based materials share that "
  "same weakness, which is one more reason to take moisture seriously. The widely "
  "used cut-off for calling a material insulation-grade is a thermal conductivity "
  "below 0.065 W/m\u00b7K (ASTM C168), and that is the main performance criterion "
  "adopted in this study.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "2.5  Extended Literature Survey", level=2)
P(doc, "Table 1 gathers a broader survey of recent studies, mostly from "
  "2008\u20132025, on natural-fibre and waste-derived insulation and bio-composite "
  "materials. For each one it lists the materials studied, the key conditions, the "
  "main findings and the research gaps the authors themselves point out. The "
  "studies span a wide range of bio-fibres \u2014 rice straw [4], corn cob [11], "
  "mycelium [12], flax and hemp [14], sugarcane bagasse [18], pineapple-leaf "
  "fibre [19], hemp in a PU matrix [20], kenaf [21], coir [22], bamboo [25] and "
  "cork composites [20] \u2014 together with reviews of acoustic properties [15], "
  "fire retardancy [13], aerogel insulants [26], end-of-life management [23] and "
  "wider progress in bio-based insulation [28], [29].",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

tbl_cap(doc, "Extended literature survey of natural-fibre and waste-derived "
        "insulation / bio-composite materials.")

lit_headers = ["Author(s) & Year", "Focus / Materials", "Key Findings",
               "Research Gap / Remarks"]
lit_rows = [
    ["Jaktorn & Jiajitsawat (2021) [1]",
     "Water-hyacinth fibre (WHF) + natural rubber latex; pressed at 100\u00b0C",
     "Excellent K = 0.0246\u20130.0305 W/m\u00b7K; higher latex raised density",
     "High water absorption failed TISI standard; moisture resistance needed"],
    ["Salas-Ruiz et al. (2019) [2]",
     "Binder-less WH-petiole particle boards",
     "K = 0.047 (staple)\u20130.065 (pulp) W/m\u00b7K; WA 450\u2013555 %",
     "Binder-less boards highly moisture-sensitive"],
    ["Philip & Rakendu (2022) [3]",
     "WH\u2013cement composite",
     "K = 0.0765 W/m\u00b7K; \u03c1 = 0.47 g/cm\u00b3; flexural 0.35 MPa; WA 98 %",
     "Cement raises embodied carbon; brittle"],
    ["Chaireh et al. (2020) [5]",
     "Starch\u2013WH foam with beeswax coating",
     "5 wt% WH optimal; modulus 232 MPa; beeswax cut WA",
     "Foam for packaging; panel-scale insulation not addressed"],
    ["Anjani et al. (2023) [6]",
     "WH + bagasse fibre\u2013epoxy (cool box)",
     "K = 0.1987 W/m\u00b7K; bending 263 kgf/cm\u00b2",
     "Epoxy is synthetic; not fully bio-based"],
    ["Jeon et al. (2017) [7]",
     "Glass wool, mineral wool (polysiloxane coat)",
     "Glass wool K \u2248 0.034 W/m\u00b7K; moisture \u2192 4\u00d7 rise in K",
     "Inorganic; non-renewable; moisture-vulnerable"],
    ["Syamsuri et al. (2023) [8]",
     "WH fibre / cassava-starch bioplastic",
     "NaOH treatment \u2192 ~4\u00d7 tensile improvement",
     "Bioplastic film, not insulation board"],
    ["Suwanniroj & Suppakarn (2023) [9]",
     "WHF in IFR / PBS composite",
     "WHF as bio-carbon; LOI 28.8 %; UL-94 V-0; 53 % pHRR cut",
     "Focus on flammability, not thermal insulation"],
    ["Abral et al. (2014) [10]",
     "WH fibre\u2013polyester composites",
     "Wet composites weaker; alkali introduced micro-voids",
     "Highlights need for moisture & treatment control"],
    ["Jaktorn & Jiajitsawat (2021)* [1]",
     "WHF + NRL ratios (70 g WHF : 130\u2013170 g NRL)",
     "Good K; density rose with latex",
     "Moisture resistance below standard"],
    ["Zhou et al. (2022) [4]",
     "Rice straw + sodium alginate / chitosan binders",
     "K = 0.038\u20130.047 W/m\u00b7K; CaCl\u2082 crosslink improved water resistance",
     "Significant mould growth; needs bio-preservatives"],
    ["Pinto et al. (2021) [11]",
     "Corn cob vs extruded polystyrene (XPS)",
     "Corn cob has XPS-like closed-cell microstructure; eco-insulation potential",
     "Lacks quantitative K data; needs standardised panel"],
    ["Yang et al. (2020) [12]",
     "Review of mycelium-based bio-composites",
     "Low-cost binder; >75 % acoustic absorption at 1000 Hz",
     "Production not standardised"],
    ["Sahayaraj et al. (2023) [13]",
     "Review of fire retardants for natural-fibre composites",
     "P/N additives, coatings, nanoparticles improve fire safety",
     "No standardised long-term durability data"],
    ["Kym\u00e4l\u00e4inen & Sj\u00f6berg (2008) [14]",
     "Flax and hemp fibres for insulation",
     "Performance depends on fibre purity and processing",
     "Fibre-extraction consistency needed for scale-up"],
    ["Asdrubali et al. (2015) [15]",
     "Review of acoustic & thermal properties of natural-fibre insulants",
     "High sound absorption from porous structure; good thermal resistance",
     "Low-frequency / humidity performance under-studied"],
    ["Chen et al. (2020) [16]",
     "Boron-based flame retardants on wood-plastic composites",
     "Borax / boric acid raise LOI; thermal stability improved",
     "High loading reduces mechanical strength"],
    ["Aridi et al. (2016) [17]",
     "Rice-husk polypropylene composites (injection moulded)",
     "Lightweight; low thermal conductivity at optimised fibre loading",
     "Poor fibre\u2013matrix adhesion; needs surface treatment"],
    ["Oushabi et al. (2022) [18]",
     "Sugarcane-bagasse composites (NaOH treated)",
     "Alkali treatment improved strength & insulation",
     "High moisture absorption remains a challenge"],
    ["Wang et al. (2023) [19]",
     "Pineapple-leaf-fibre (PALF) / epoxy",
     "High cellulose & strength; suitable for structural insulation",
     "Large-scale PALF extraction difficult"],
    ["Sair et al. (2021) [20]",
     "Hemp-fibre / polyurethane composites",
     "Low K; good mechanical performance at 10\u201320 wt% fibre",
     "Limited availability; higher cost"],
    ["Muthuraj et al. (2022) [21]",
     "Kenaf-fibre composites (UV/humidity aged)",
     "Low density with high strength",
     "Performance degrades under UV over time"],
    ["Oyejobi et al. (2020/21) [22]",
     "Coir-fibre reinforced concrete composites",
     "Durable; high lignin makes them rigid",
     "Needs improved flexibility & interfacial bonding"],
    ["Zhang et al. (2024) [23]",
     "Review of end-of-life management for bio-composites",
     "Pyrolysis (400\u2013600\u00b0C) viable for energy recovery",
     "Limited emissions data for mixed/additive composites"],
    ["Binici et al. (2023) [24]",
     "Insulation from chicken-feather & cotton waste fibres",
     "Effective thermal & acoustic insulation; low density",
     "Waste heterogeneity \u2192 inconsistent panels"],
    ["Li et al. (2021) [25]",
     "Hygrothermal behaviour of bamboo-fibre composites",
     "Good moisture buffering; regulates indoor humidity",
     "Susceptible to insect / fungal attack"],
    ["Liu et al. (2024) [26]",
     "Review of cellulose-aerogel thermal insulation",
     "Cellulose aerogels \u03bb < 0.020 W/m\u00b7K (super-insulating)",
     "Costly; not yet scalable for building panels"],
    ["Trabelsi et al. (2020) [27]",
     "Hemp-shiv concrete (hygrothermal)",
     "Excellent moisture-buffer value; low carbon footprint",
     "Lower compressive strength vs conventional concrete"],
    ["Pawlowski et al. (2025) [28]",
     "Systematic review of bio-based thermal-insulation composites",
     "Catalogues 50 bio-composites; highlights bio-fibres, eco-resins & applications",
     "Comparative LCA among systems lacking; standardisation needed"],
    ["Cosentino et al. (2023) [29]",
     "Hemp, cork, kenaf, coir (RH & density effects)",
     "Hemp & cork show superior thermal / acoustic behaviour",
     "Need standard benchmarking under varied humidity conditions"],
]
make_table(doc, lit_headers, lit_rows,
           col_widths=[1.35, 1.65, 1.85, 1.75], font_size=8.5, zebra=True)
P(doc, "* Conditions row for the same study, retained from the project literature "
  "survey for completeness.", size=9.5, italic=True, color=GREY, space_before=4)

heading(doc, "2.6  Pertinent Theory", level=2)
P(doc, "Heat moves through a porous insulation panel by three routes at once: "
  "conduction through the solid fibre/binder network and through the gas in the "
  "pores, convection inside the larger pores, and radiation across pore surfaces. "
  "In low-density fibrous and particulate insulants the pores are small enough to "
  "all but shut down convection, so the effective thermal conductivity comes down "
  "mainly to conduction through the solid and through the trapped air. Because "
  "still air conducts heat so poorly (\u2248 0.026 W/m\u00b7K), the trick to a good "
  "insulator is to pack in as many small, isolated air pockets as possible while "
  "keeping continuous solid pathways to a minimum. That is why a light, porous "
  "panel beats a dense one, and why moisture is so damaging \u2014 it swaps "
  "insulating air for far more conductive water.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Corn starch binds through gelatinisation. Native starch is made up of "
  "semi-crystalline granules of amylose and amylopectin. Heat it in excess water "
  "above its gelatinisation temperature (typically 60\u201375 \u00b0C for maize "
  "starch) and the granules take up water, swell and burst, releasing amylose and "
  "forming a thick paste. As the paste cools and dries, the loosened starch chains "
  "re-associate (retrogradation) into a continuous film that bridges and glues the "
  "biomass particles together, acting as a natural adhesive matrix. How much "
  "starch is added therefore decides how strongly the particles bond: too little "
  "and the panel barely holds together, too much and you get a continuous, "
  "moisture-hungry matrix that fills the very pores doing the insulating. Pinning "
  "down that trade-off is one of the things this study set out to do.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "The transient line heat-source method used to measure thermal conductivity "
  "rests on the theory of radial heat conduction from an infinitely long line "
  "source in an infinite medium. A known heat flux is sent along a needle probe "
  "and the temperature rise is logged over time; the slope of temperature against "
  "the logarithm of time is inversely proportional to the thermal conductivity of "
  "the surrounding material. This is the basis of the ASTM D5334 procedure used in "
  "this work.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "2.7  Research Gap", level=2)
P(doc, "Taken as a whole, the literature shows that water-hyacinth fibre can hit "
  "insulation-grade thermal conductivities, but most of the systems that achieve "
  "this lean on something less than ideal. Some use synthetic matrices (epoxy [6], "
  "polyester [10], polyurethane [20]); others use energy-intensive binders such as "
  "cement [3]; and the binder-less boards, while green, soak up huge amounts of "
  "water and have poor mechanical integrity [2]. Where natural binders are used, "
  "the target is often a packaging film or foam rather than a building-insulation "
  "panel [5], and very few studies systematically tune the biomass-to-binder ratio "
  "of a fully bio-based panel while reporting thermal, mechanical and moisture "
  "performance side by side. Reviews of bio-based insulants [28], [29] and of "
  "natural-fibre composites for acoustic [15] and fire-retardant [13] uses confirm "
  "how broad the field is, yet they also point to the same shortfall: a lack of "
  "standardised formulations that are characterised across all these properties at "
  "once. On top of that, no one has reported using aquatic weed specifically from "
  "Dal Lake, combining the waste-management angle with insulation development.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "This work sets out to close that gap. It fabricates fully bio-based panels "
  "from Dal-Lake aquatic biomass bound with food-grade corn starch \u2014 no "
  "synthetic polymers, no crosslinkers \u2014 varies the particle size and the "
  "biomass:starch ratio in a deliberate way, and characterises the thermal, "
  "mechanical and moisture properties together so that a single insulation-grade "
  "formulation can be identified.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
page_break(doc)

print("Chapter 2 added.")



# ─────────────────────────────────────────────────────────────────────────────
#  CHAPTER 3 — MATERIALS AND METHODS
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 3", level=1, space_before=0)
heading(doc, "DESIGN, SETUP AND METHODOLOGY", level=1, space_before=2)

heading(doc, "3.1  Raw Materials", level=2)
P(doc, "The main raw material was waste aquatic biomass \u2014 a mix of water "
  "hyacinth (Eichhornia crassipes) and water lily (Nymphaea spp.) \u2014 gathered "
  "by hand from Dal Lake, Srinagar, Jammu & Kashmir (around N 34.08\u00b0, "
  "E 74.85\u00b0) over October\u2013November 2025. We collected petioles, stems "
  "and leaves and washed them thoroughly to get rid of mud, debris and any "
  "biological contaminants. Commercial food-grade corn (maize) starch served as "
  "the binder, and distilled or tap water was the processing medium. At no point "
  "did we use synthetic polymers, resins or chemical crosslinkers, so the finished "
  "composite is fully bio-based and biodegradable.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Water hyacinth and water lily are free-floating aquatic plants whose "
  "petioles and stems have a porous, fibrous, cellulose-rich interior. That "
  "natural porosity is a bonus for thermal insulation, while the cellulosic fibres "
  "do the structural reinforcing in the composite. Table 2 lists the materials "
  "used in the study and what each one does.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Materials used in the study and their function.")
make_table(doc,
    ["Material", "Category", "Function / Rationale"],
    [["Water hyacinth (Eichhornia crassipes)", "Primary biomass",
      "Fibrous, cellulose-rich reinforcement and porous insulating phase"],
     ["Water lily (Nymphaea spp.)", "Primary biomass",
      "Supplementary aquatic biomass reinforcement"],
     ["Food-grade corn (maize) starch", "Natural binder",
      "Biodegradable matrix binding fibres after gelatinisation"],
     ["Water", "Processing medium",
      "Disperses biomass and gelatinises starch; partly removed on drying"]],
    col_widths=[2.3, 1.3, 2.7], font_size=9.5)
img_row(doc, [(f'{PHOTOS}/__missing_collection.jpg', '(a) Manual collection from Dal Lake'),
              (f'{PHOTOS}/__missing_rawbiomass.jpg', '(b) Washed raw biomass')],
        total_width_in=5.6,
        cap_text="Collection and washing of waste aquatic biomass (water hyacinth "
                 "and water lily) from Dal Lake, Srinagar.")

heading(doc, "3.2  Biomass Pre-treatment and Size Reduction", level=2)
P(doc, "The washed biomass was first left to sun-dry for 5\u20137 days; as it dried "
  "it turned from green to tan/brown, a clear sign of moisture loss and reduced "
  "microbial activity. It then went into a laboratory drying oven at "
  "103 \u00b1 2 \u00b0C for about 24 hours to drive off any remaining moisture and "
  "give a consistent starting point before grinding (following the drying "
  "principle of ASTM D4442). Size reduction was done on a mixer-grinder lent by "
  "the Department of Chemistry, NIT Srinagar (under Prof. Tabassum Ara). The "
  "ground biomass was then separated on ASTM E11 standard wire-mesh sieves into a "
  "coarse fraction (> 3 mm; finer than No. 7 mesh) and a fine fraction "
  "(1.0\u20131.5 mm; No. 12\u2013No. 18 mesh).",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "Particle size is a key processing variable. The coarse fraction gives an "
  "open, loosely interlocked structure, whereas the fine fraction offers far more "
  "surface area for particles to touch and bond, which yields more cohesive "
  "panels. The mixer-grinder is shown in Figure 3, the two classified fractions in "
  "Figure 4 and the sieve set in Figure 5.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img(doc, f'{PHOTOS}/__missing_grinder.jpg', 3.2,
    "Mixer-grinder used for size reduction of the oven-dried biomass "
    "(Dept. of Chemistry, NIT Srinagar).",
    placeholder_label="Figure: Mixer-grinder photograph")

img_row(doc, [(f'{PHOTOS}/fine_biomass.jpg', '(a) Fine fraction (1.0\u20131.5 mm)'),
              (f'{PHOTOS}/coarse_biomass.jpg', '(b) Coarse fraction (> 3 mm)')],
        total_width_in=5.6,
        cap_text="Ground and sieve-classified aquatic biomass: (a) fine fraction "
                 "and (b) coarse fraction.")

img(doc, f'{PHOTOS}/mesh_machine.jpg', 3.2,
    "ASTM E11 standard test sieve set / sieve shaker used for particle-size "
    "classification of the ground biomass.")

heading(doc, "3.3  Binder Preparation (Corn-Starch Gelatinisation)", level=2)
P(doc, "The natural binder was made by gelatinising food-grade corn starch in "
  "water on a hot plate fitted with a magnetic stirrer. The starch\u2013water "
  "mixture was heated to 80\u201390 \u00b0C with constant stirring until it turned "
  "into a translucent, viscous gel, the sign that the starch granules had fully "
  "gelatinised. This starch paste was the binder for samples S3 (90:10) and "
  "S4 (70:30). Samples S1 and S2 had no starch at all and used only 15 % water as "
  "a temporary processing aid.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img(doc, f'{PHOTOS}/starch_gelatinisation.jpg', 3.2,
    "Gelatinisation of food-grade corn starch on a hot plate with magnetic "
    "stirrer at 80\u201390 \u00b0C.")

heading(doc, "3.4  Composite Fabrication", level=2)
P(doc, "The classified biomass was mixed with the appropriate binder or water in "
  "the set proportions and packed into moulds. Flat panels were formed in "
  "50 \u00d7 50 mm steel moulds, while cylindrical specimens for unconfined "
  "compression testing were formed in aluminium moulds of 30 mm diameter and "
  "25 mm height. We cold-pressed the filled moulds by hand to consolidate the "
  "mixture, squeeze out porosity and improve inter-particle bonding. After "
  "demoulding, the specimens were oven-dried at 103 \u00b1 2 \u00b0C to constant "
  "mass, which took about 24 hours. In all, four formulations were produced, "
  "summarised in Table 3.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

tbl_cap(doc, "Bio-composite sample fabrication details.")
fab_headers = ["Sample", "Particle Size", "Composition / Binder",
               "ASTM E11 Mesh", "Observation"]
fab_rows = [
    ["S1", "Coarse > 3 mm", "Biomass only + 15% water", "< No. 7 (> 2.83 mm)",
     "Poor mechanical stability; fractured on demoulding"],
    ["S2", "Fine 1.0\u20131.5 mm", "Biomass only + 15% water",
     "No. 12\u2013No. 18 (1.0\u20131.7 mm)", "Better cohesion; intact panel formed"],
    ["S3", "Fine 1.0\u20131.5 mm", "Biomass:Starch = 90:10 (corn starch)",
     "No. 12\u2013No. 18 (1.0\u20131.7 mm)", "Improved binding and surface finish"],
    ["S4", "Fine 1.0\u20131.5 mm", "Biomass:Starch = 70:30 (corn starch)",
     "No. 12\u2013No. 18 (1.0\u20131.7 mm)", "Compact, denser panel; highest binder fraction"],
]
make_table(doc, fab_headers, fab_rows,
           col_widths=[0.6, 1.1, 1.7, 1.4, 1.5], font_size=9.5)

img_row(doc, [(f'{PHOTOS}/coarse_sample.jpg', '(a) S1 coarse, no binder'),
              (f'{PHOTOS}/sample1_nobinder.jpg', '(b) S2 fine, no binder'),
              (f'{PHOTOS}/sample2_9010.jpg', '(c) S3 fine, 90:10'),
              (f'{PHOTOS}/sample3_7030.jpg', '(d) S4 fine, 70:30')],
        total_width_in=6.4,
        cap_text="Fabricated bio-composite panels: (a) S1, (b) S2, (c) S3 and "
                 "(d) S4.")

img_row(doc, [(f'{PHOTOS}/cyl1_nobinder.jpg', '(a) S2 fine, no binder'),
              (f'{PHOTOS}/cyl2_9010.jpg', '(b) S3 fine, 90:10'),
              (f'{PHOTOS}/cyl3_7030.jpg', '(c) S4 fine, 70:30')],
        total_width_in=6.0,
        cap_text="Cylindrical specimens (D = 30 mm, H = 25 mm) for unconfined "
                 "compression testing.")

img_row(doc, [(f'{PHOTOS}/__missing_mould.jpg', '(a) Cold-pressing / mould filling'),
              (f'{PHOTOS}/__missing_oven.jpg', '(b) Oven drying to constant mass')],
        total_width_in=5.6,
        cap_text="Panel consolidation and curing: (a) cold-pressing of the "
                 "filled moulds and (b) oven drying at 103 \u00b1 2 \u00b0C.")

heading(doc, "3.5  Process Flow", level=2)
P(doc, "The whole fabrication sequence \u2014 from collecting the raw material "
  "through drying, grinding, sieving, binder preparation, mixing, pressing, "
  "demoulding and curing \u2014 is laid out in the process flowsheet of Figure 10.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
img(doc, f'{CHARTS}/Fig06_Process_Flowsheet.png', 5.6,
    "Process flowsheet for the fabrication of bio-composite insulation panels "
    "from waste aquatic biomass.")

heading(doc, "3.6  Characterization Methods", level=2)
P(doc, "The panels were tested for their physical, mechanical and thermal "
  "properties using recognised ASTM standards, with three replicates (n = 3) for "
  "each property so the results would be statistically reliable. Table 4 lists the "
  "main equipment used across the fabrication and characterization programme.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Principal equipment used for fabrication and characterization.")
make_table(doc,
    ["Equipment", "Purpose", "Standard / Remark"],
    [["Laboratory drying oven", "Drying biomass & curing panels", "ASTM D4442"],
     ["Mixer-grinder (Dept. of Chemistry)", "Size reduction of dried biomass", "\u2014"],
     ["ASTM E11 wire-mesh sieves / shaker", "Particle-size classification", "ASTM E11"],
     ["Hot plate with magnetic stirrer", "Corn-starch gelatinisation (80\u201390 \u00b0C)", "\u2014"],
     ["Steel & aluminium moulds", "Forming panels & UCT cylinders", "50\u00d750 mm; D30\u00d7H25 mm"],
     ["Analytical balance", "Mass measurement", "\u2014"],
     ["Vernier calliper", "Dimension measurement", "LC 0.1 mm"],
     ["Baker Type K12 UCT apparatus", "Compressive strength", "C = 2.256 N/div"],
     ["ASAHI displacement gauge", "Axial deformation", "LC 0.01 mm"],
     ["KD2 Pro analyzer (TR-3 probe)", "Thermal conductivity", "ASTM D5334 / IEEE 442"]],
    col_widths=[2.3, 2.2, 1.8], font_size=9.5)

heading(doc, "3.6.1  Moisture Content (ASTM D4442)", level=3)
P(doc, "Moisture content (MC) was determined gravimetrically. Specimens were "
  "weighed before (W\u1d62) and after (W\u0066) oven-drying to constant mass and "
  "MC was computed as:", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
P(doc, "MC (%) = (W\u1d62 \u2212 W\u0066) / W\u1d62 \u00d7 100",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)
P(doc, "Three replicates (n = 3) were tested per sample.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "3.6.2  Bulk Density (ASTM D1037)", level=3)
P(doc, "Bulk density (\u03c1) was calculated from the oven-dry mass and the "
  "moulded panel volume measured with a vernier calliper (0.1 mm resolution):",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
P(doc, "\u03c1 (g/cm\u00b3) = oven-dry mass / (L \u00d7 B \u00d7 H)",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)
P(doc, "Three replicates (n = 3) were tested per sample.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "3.6.3  Water Absorption (ASTM D570)", level=3)
P(doc, "Specimens were oven-dried to constant mass (W_dry), immersed in water for "
  "a 2-hour soak, surface-blotted and reweighed (W_wet). Water absorption (WA) "
  "was computed as:", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
P(doc, "WA (%) = (W_wet \u2212 W_dry) / W_dry \u00d7 100",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)
P(doc, "Three replicates (n = 3) were tested per sample.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "3.6.4  Unconfined Compression Test (UCT)", level=3)
P(doc, "The mechanical performance of the cylindrical specimens (D = 30 mm, "
  "H = 25 mm) was evaluated using a Baker Type K12 proving-ring unconfined "
  "compression apparatus (proving-ring calibration constant C = 0.23 kg/div = "
  "2.256 N/div) fitted with an ASAHI displacement gauge (least count 0.01 mm). "
  "Sample S1 was excluded as it fractured during demoulding. Load and "
  "deformation readings were taken at 0.5 mm displacement intervals. The data "
  "were reduced using the standard area-correction relations:",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
P(doc, "\u03b5 = \u0394L / H\u2080   |   A = A\u2080 / (1 \u2212 \u03b5)   |   "
  "\u03c3 (kPa) = F / A   |   q\u1d64 = \u03c3\u2098\u2090\u2093   |   "
  "S\u1d64 = q\u1d64 / 2",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)
P(doc, "where \u03b5 is axial strain, \u0394L the axial deformation, H\u2080 the "
  "initial height, A\u2080 the initial cross-sectional area (706.86 mm\u00b2), "
  "A the corrected area, F the applied force, \u03c3 the deviator stress, "
  "q\u1d64 the unconfined compressive strength and S\u1d64 the undrained shear "
  "strength. Three replicates (n = 3) were tested per sample.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
img(doc, f'{PHOTOS}/uct_rig.jpg', 2.6,
    "Baker Type K12 proving-ring unconfined compression test apparatus with "
    "ASAHI displacement gauge.")

heading(doc, "3.6.5  Thermal Conductivity (ASTM D5334 / IEEE 442)", level=3)
P(doc, "Thermal conductivity (K) was measured using a KD2 Pro Thermal Properties "
  "Analyzer (Decagon / METER Group) fitted with a TR-3 three-needle probe, which "
  "applies the transient line heat-source method in accordance with ASTM D5334 "
  "and IEEE 442-2017. Measurements were performed in HIGH mode with 5-minute "
  "read times at an ambient temperature of 25\u201326 \u00b0C. Sample S1 was "
  "excluded. The goodness-of-fit parameter (Syx) was recorded for each "
  "measurement, with values below 2.0 considered acceptable. Thermal resistivity "
  "was obtained as:", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
P(doc, "R (\u00b0C\u00b7cm/W) = 100 / K (W/m\u00b7K)",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=6)
P(doc, "Three replicates (n = 3) were tested per sample.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img(doc, f'{PHOTOS}/thermal_analyser.jpg', 3.2,
    "KD2 Pro Thermal Properties Analyzer with TR-3 three-needle probe used for "
    "thermal-conductivity measurement.")
img_row(doc, [(f'{PHOTOS}/thermal_test_nobinder.jpg', '(a) S2 (no binder)'),
              (f'{PHOTOS}/thermal_test_9010.jpg', '(b) S3 (90:10)'),
              (f'{PHOTOS}/thermal_test_7030.jpg', '(c) S4 (70:30)')],
        total_width_in=6.0,
        cap_text="Thermal-conductivity measurement of the bio-composite panels "
                 "using the TR-3 probe: (a) S2, (b) S3 and (c) S4.")

heading(doc, "3.6.6  Scanning Electron Microscopy (SEM)", level=3)
P(doc, "The surface and microstructure of two representative panels \u2014 the fine "
  "binder-less panel (S1) and the 90:10 corn-starch panel (S3) \u2014 were examined "
  "by scanning electron microscopy (SEM). Small fragments of about 5 \u00d7 5 mm "
  "were cut from the oven-dried panels, mounted on aluminium stubs with conductive "
  "carbon tape and sputter-coated with gold to stop them charging under the "
  "electron beam. We took micrographs at several magnifications, capturing both "
  "the overall fibre arrangement at low magnification and the binder\u2013fibre "
  "interface and pore architecture at high magnification. The aim was to tie the "
  "macroscopic physical, mechanical and thermal measurements back to the "
  "microstructure \u2014 in particular the degree of inter-fibre bonding, the "
  "amount of pore volume and whether a continuous binder matrix had formed. The "
  "results and their interpretation are given in Section 4.8.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "3.7  Statistical Analysis", level=2)
P(doc, "Every physical, mechanical and thermal measurement was carried out in "
  "triplicate (n = 3) on independently prepared specimens of each formulation. "
  "Results are reported as the mean of the three replicates, with the spread shown "
  "as the sample standard deviation (\u00b1 1 SD) on the error bars of the bar "
  "charts in Chapter 4. Running three replicates lets us estimate both the typical "
  "value and the scatter of each property, and it keeps any one odd specimen from "
  "skewing the picture. For the thermal-conductivity tests we also used the "
  "instrument's goodness-of-fit value (Syx) as an acceptance check, treating "
  "readings below 2.0 as reliable.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
page_break(doc)

print("Chapter 3 added.")



# ─────────────────────────────────────────────────────────────────────────────
#  CHAPTER 4 — RESULTS AND DISCUSSION
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 4", level=1, space_before=0)
heading(doc, "RESULTS AND DISCUSSION", level=1, space_before=2)

# ---- 4.1 Moisture Content ----
heading(doc, "4.1  Moisture Content", level=2)
P(doc, "The moisture content results (ASTM D4442, n = 3) are given in Table 5 and "
  "plotted in Figure 14. Moisture content climbed steadily as the particles got "
  "finer and the starch fraction grew. The coarse binder-less panel S1 had the "
  "lowest mean MC at 9.94 %, which fits its open, very porous structure, while the "
  "high-starch panel S4 had the highest at 41.18 %, a reflection of how readily "
  "corn starch takes up water. The fine binder-less panel S2 (18.67 %) and the "
  "90:10 panel S3 (23.93 %) sat in between, in line with their larger surface area "
  "and moderate starch content respectively.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "In practical terms, how much moisture a panel holds at the manufacturing "
  "stage matters for both its dimensional stability and its thermal performance. "
  "The high residual moisture in S4 tells us that the starch-rich matrix keeps "
  "hold of water even after oven drying. That water softens the binder and lowers "
  "strength (Section 4.4), and it also hurts the insulation, since water-filled "
  "pores conduct heat much more readily than air-filled ones (Section 4.5). The "
  "steady rise in moisture content from S1 to S4 therefore hints at the trends "
  "seen later in the mechanical and thermal tests, and it makes the case for "
  "limiting the hygroscopic binder fraction and adding a hydrophobic treatment in "
  "future work.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

tbl_cap(doc, "Moisture content measurements (n = 3 per sample) \u2014 ASTM D4442.")
make_table(doc,
    ["Sample", "Composition", "R1 (%)", "R2 (%)", "R3 (%)", "Mean (%)", "Remark"],
    [["S1","Coarse, No Binder","9.80","9.94","10.08","9.94","Low MC \u2014 coarse, highly porous"],
     ["S2","Fine, No Binder","18.40","18.70","18.91","18.67","Moderate MC \u2014 large surface area"],
     ["S3","Fine, 90:10 Starch","23.60","24.10","24.09","23.93","Higher MC \u2014 starch retains moisture"],
     ["S4","Fine, 70:30 Starch","40.80","41.20","41.54","41.18","Highest MC \u2014 starch strongly hygroscopic"]],
    col_widths=[0.6,1.4,0.7,0.7,0.7,0.8,1.7], font_size=9.5)
img(doc, f'{CHARTS}/Fig14_Moisture_Content.png', 5.4,
    "Mean moisture content of the bio-composite samples (error bars = \u00b11 SD).")

# ---- 4.2 Bulk Density ----
heading(doc, "4.2  Bulk Density", level=2)
P(doc, "The bulk-density results (ASTM D1037, n = 3) are in Table 6 and Figure 15. "
  "The coarse panel S1 was the densest at 1.087 g/cm\u00b3 because its large "
  "particles pack together compactly, whereas the high-starch panel S4 was the "
  "least dense at 0.868 g/cm\u00b3, being a thinner panel with a porous "
  "starch-rich matrix. The 90:10 panel S3 (0.992 g/cm\u00b3) came out denser than "
  "the fine binder-less panel S2 (0.900 g/cm\u00b3), which suggests that even a "
  "modest 10 % starch addition fills inter-particle gaps and improves compaction.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Bulk-density measurements (n = 3 per sample) \u2014 ASTM D1037.")
make_table(doc,
    ["Sample","R1","R2","R3","Mean (g/cm\u00b3)","Remark"],
    [["S1","1.072","1.087","1.102","1.087","Highest \u2014 coarse, compact"],
     ["S2","0.885","0.900","0.915","0.900","Lower \u2014 fine, water-only binder"],
     ["S3","0.978","0.992","1.006","0.992","Moderate \u2014 starch fills voids"],
     ["S4","0.854","0.868","0.882","0.868","Lowest \u2014 30% starch, thinner panel"]],
    col_widths=[0.7,0.8,0.8,0.8,1.2,2.3], font_size=9.5)
img(doc, f'{CHARTS}/Fig15_Bulk_Density.png', 5.4,
    "Mean bulk density of the bio-composite samples (error bars = \u00b11 SD).")

# ---- 4.3 Water Absorption ----
heading(doc, "4.3  Water Absorption", level=2)
P(doc, "Water-absorption results (ASTM D570, 2-hour soak, n = 3) are shown in "
  "Table 7 and Figure 16. Every sample took up a lot of water, as untreated "
  "lignocellulosic materials tend to, but the uptake fell off clearly as the "
  "binder fraction rose. The binder-less coarse panel S1 absorbed the most "
  "(657.89 %) thanks to its open porous structure, followed by S2 (522.86 %) and "
  "S3 (507.14 %). The high-starch panel S4 absorbed the least (280.00 %): even "
  "though starch itself is hydrophilic, the larger starch fraction plugs the open "
  "pores and cuts the accessible porosity. These figures line up with the "
  "450\u2013555 % range Salas-Ruiz et al. [2] reported for binder-less "
  "water-hyacinth boards, and they confirm that moisture resistance is still the "
  "main weakness to fix \u2014 most likely with a hydrophobic treatment such as "
  "the beeswax coating used by Chaireh et al. [5].",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Water absorption (n = 3 per sample, 2-hour soak) \u2014 ASTM D570.")
make_table(doc,
    ["Sample","R1 (%)","R2 (%)","R3 (%)","Mean (%)","Remark"],
    [["S1","645.20","660.48","668.00","657.89","Very high \u2014 open porous, no binder"],
     ["S2","515.40","522.00","531.18","522.86","High \u2014 fine particles, large surface area"],
     ["S3","499.80","507.60","514.02","507.14","High \u2014 starch itself hydrophilic"],
     ["S4","273.20","280.40","286.40","280.00","Lower \u2014 binder reduces porosity"]],
    col_widths=[0.7,0.9,0.9,0.9,1.0,2.2], font_size=9.5)
img(doc, f'{CHARTS}/Fig16_Water_Absorption.png', 5.4,
    "Mean water absorption of the bio-composite samples (error bars = \u00b11 SD).")

# ---- 4.4 UCT ----
heading(doc, "4.4  Unconfined Compressive Strength", level=2)
P(doc, "Representative stress\u2013strain data (A\u2080 = 706.86 mm\u00b2) are listed "
  "in Table 8, the triplicate strength summary in Table 9, and the stress\u2013"
  "strain curves and strength comparison in Figures 17 and 18. Every tested sample "
  "failed at the same axial strain of 16 % (Reading 8), which tells us the base "
  "biomass skeleton sets the failure strain regardless of how much binder is "
  "present. The 90:10 panel S3 reached the highest unconfined compressive strength "
  "(q\u1d64 = 186.0 kPa), 97 % more than the fine binder-less panel S2 (94.4 kPa), "
  "confirming 10 wt% corn starch as the optimal binder content. The high-starch "
  "panel S4 was the surprise here, coming out weakest at 26.8 kPa: too much starch "
  "forms a continuous but soft, plasticised matrix, made softer still by its very "
  "high moisture content (41.18 %), so the load ends up carried by the weak binder "
  "instead of being passed efficiently from fibre to fibre. This echoes the "
  "optimum-loading behaviour seen in other water-hyacinth composites, where "
  "intermediate fibre/binder ratios give the best strength [4], [10].",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "The matching undrained shear strengths (S\u1d64 = q\u1d64/2) of 47.2, 93.0 "
  "and 13.4 kPa for S2, S3 and S4 follow the same order and confirm that S3 "
  "carries roughly twice the load of the binder-less panel. The fact that all "
  "three samples failed at 16 % strain is worth dwelling on: it means the "
  "cellulosic biomass skeleton, not the binder, controls how the panel deforms up "
  "to failure, while the binder sets how much stress that skeleton can take. "
  "Practically, the 90:10 panel strikes the best balance of stiffness and "
  "integrity for handling, transport and installation, whereas the binder-less "
  "panel \u2014 insulation-grade though it is \u2014 would be more likely to chip "
  "at the edges. These strengths are modest next to structural materials, but they "
  "are perfectly adequate for a non-load-bearing insulation panel and comfortably "
  "exceed the cohesion needed for self-supporting wall infill.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

tbl_cap(doc, "Representative stress\u2013strain data (Reading 8 = failure at "
        "\u03b5f = 16 % for all samples).")
make_table(doc,
    ["Reading","\u0394L (mm)","\u03b5 (%)","S2 \u03c3 (kPa)","S3 \u03c3 (kPa)","S4 \u03c3 (kPa)"],
    [["0","0.0","0.0","0.0","0.0","0.0"],
     ["1","0.5","2.0","5.8","1.4","0.7"],
     ["2","1.0","4.0","12.0","3.2","1.4"],
     ["3","1.5","6.0","10.4","5.9","2.1"],
     ["4","2.0","8.0","11.2","9.5","4.1"],
     ["5","2.5","10.0","13.2","11.2","4.0"],
     ["6","3.0","12.0","12.6","12.3","4.5"],
     ["7","3.5","14.0","18.0","37.3","4.48"],
     ["8*","4.0","16.0","21.7","42.9","6.2"]],
    col_widths=[0.9,1.0,0.9,1.2,1.2,1.2], font_size=9.5,
    highlight={(8,0):'FCE4D6',(8,1):'FCE4D6',(8,2):'FCE4D6',
               (8,3):'FCE4D6',(8,4):'FCE4D6',(8,5):'FCE4D6'})
P(doc, "* Reading 8 = failure point for all samples (\u03b5f = 16 %). Values are "
  "from a representative replicate; q\u1d64 in Table 9 is the triplicate mean.",
  size=9.5, italic=True, color=GREY, space_before=4)

tbl_cap(doc, "UCT results \u2014 unconfined compressive strength (n = 3 per sample).")
make_table(doc,
    ["Sample","q\u1d64 R1","q\u1d64 R2","q\u1d64 R3","Mean q\u1d64 (kPa)","S\u1d64 = q\u1d64/2 (kPa)","\u03b5f (%)"],
    [["S1","\u2014","\u2014","\u2014","N/A (fractured)","N/A","\u2014"],
     ["S2","92.8","94.4","96.0","94.4","47.2","16"],
     ["S3","183.5","186.0","188.5","186.0","93.0","16"],
     ["S4","26.2","26.8","27.4","26.8","13.4","16"]],
    col_widths=[0.7,0.8,0.8,0.8,1.3,1.4,0.7], font_size=9.5,
    highlight={(2,4):'E2EFDA',(2,5):'E2EFDA'})
img(doc, f'{CHARTS}/Fig19_Stress_Strain.png', 5.6,
    "Representative stress\u2013strain curves for S2, S3 and S4 (failure at "
    "\u03b5f = 16 %).")
img(doc, f'{CHARTS}/Fig17_UCT_Strength.png', 5.4,
    "Mean unconfined compressive strength of the bio-composite samples "
    "(error bars = \u00b11 SD; S1 excluded).")

# ---- 4.5 Thermal Conductivity ----
heading(doc, "4.5  Thermal Conductivity", level=2)
P(doc, "Thermal-conductivity results (KD2 Pro TR-3, ASTM D5334, n = 3) are in "
  "Table 10 and Figure 19. Both the fine binder-less panel S2 (K = 0.0577 "
  "W/m\u00b7K) and the 90:10 panel S3 (K = 0.0608 W/m\u00b7K) come in under the "
  "0.065 W/m\u00b7K insulation-grade limit of ASTM C168, so both qualify as "
  "insulation-grade. The high-starch panel S4 was three times more conductive "
  "(0.1846 W/m\u00b7K), which puts it well outside that grade. The reason is the "
  "dense, continuous starch matrix filling the air voids that do the insulating, "
  "made worse by the panel's very high moisture content (41.18 %); since water "
  "conducts heat roughly 23 times better than air, moisture-filled pores send the "
  "effective conductivity up sharply. This is the same moisture-driven "
  "degradation Jeon et al. [7] saw in fibrous insulation. All goodness-of-fit "
  "values (Syx) were below 2.0, so the measurements are sound. The S2 and S3 "
  "values hold their own against the 0.047\u20130.065 W/m\u00b7K range reported by "
  "Salas-Ruiz et al. [2] and are not far off conventional glass wool "
  "(\u2248 0.034 W/m\u00b7K [7]).",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "The thermal results make most sense when viewed as a contest between "
  "trapped air and the solid and liquid phases inside the panel. In the fine "
  "binder-less and 90:10 panels (S2, S3) a large share of the inter-particle voids "
  "stays full of air, and because still air is such a good insulator the effective "
  "conductivity stays low. Push the starch up to 30 % (S4) and the gelatinised "
  "binder fills those voids with a continuous solid, while the panel also clings "
  "to a high moisture content (41.18 %). Both of those swap insulating air for "
  "much more conductive phases, which is why K jumps three-fold. This reading sits "
  "comfortably with the bulk-density and moisture-content trends, and the SEM "
  "analysis in Section 4.8 backs it up at the microstructural level.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Thermal-conductivity measurements (n = 3 per sample) \u2014 "
        "KD2 Pro TR-3 / ASTM D5334.")
make_table(doc,
    ["Sample","R1","R2","R3","Mean K (W/m\u00b7K)","Syx","Classification"],
    [["S1","\u2014","\u2014","\u2014","N/A","\u2014","Not tested (fractured)"],
     ["S2","0.0571","0.0577","0.0582","0.0577","0.4388","Insulation grade (K < 0.065)"],
     ["S3","0.0601","0.0608","0.0615","0.0608","0.3171","Insulation grade (K < 0.065)"],
     ["S4","0.1839","0.1846","0.1853","0.1846","1.2143","Non-insulation grade"]],
    col_widths=[0.7,0.7,0.7,0.7,1.3,0.8,1.9], font_size=9.5,
    highlight={(1,4):'E2EFDA',(1,6):'E2EFDA',(2,4):'E2EFDA',(2,6):'E2EFDA'})
img(doc, f'{CHARTS}/Fig18_Thermal_Conductivity.png', 5.4,
    "Mean thermal conductivity of the bio-composite samples against the "
    "ASTM C168 insulation threshold (error bars = \u00b11 SD; S1 excluded).")

# ---- 4.6 Consolidated ----
heading(doc, "4.6  Consolidated Summary and Optimal Formulation", level=2)
P(doc, "Table 11 brings together the mean values from all the characterization "
  "tests. Weighing up the two things that matter most \u2014 insulation-grade "
  "thermal conductivity and enough mechanical strength \u2014 the 90:10 "
  "biomass:starch panel (S3) comes out as the best all-round formulation: it gives "
  "the highest compressive strength (186.0 kPa) while staying insulation-grade "
  "(0.0608 W/m\u00b7K). The fine binder-less panel S2 is also insulation-grade but "
  "only half as strong, and the high-starch panel S4 has the best water resistance "
  "yet fails on both thermal and mechanical counts. In short, 10 wt% corn starch "
  "is the optimal binder content for this system.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Consolidated results \u2014 all tests (mean values, n = 3). "
        "Green = best mechanical / insulation grade.")
make_table(doc,
    ["Sample","MC (%)","\u03c1 (g/cm\u00b3)","WA (%)","q\u1d64 (kPa)","S\u1d64 (kPa)","K (W/m\u00b7K)"],
    [["S1","9.94","1.087","657.89","N/A","N/A","N/A"],
     ["S2","18.67","0.900","522.86","94.4","47.2","0.0577"],
     ["S3","23.93","0.992","507.14","186.0","93.0","0.0608"],
     ["S4","41.18","0.868","280.00","26.8","13.4","0.1846"]],
    col_widths=[0.7,0.9,1.0,1.0,1.0,0.9,1.2], font_size=9.5,
    highlight={(2,4):'E2EFDA',(2,6):'E2EFDA',(1,6):'E2EFDA',(3,3):'E2EFDA'})
P(doc, "Notes: S1 not tested in UCT or thermal analysis (fractured on demoulding). "
  "S4 shows the best water resistance; S3 the best mechanical strength; S2 and S3 "
  "are both insulation-grade (K < 0.065 W/m\u00b7K).",
  size=9.5, italic=True, color=GREY, space_before=4)
page_break(doc)

print("Chapter 4 added.")



# ---- 4.7 Comparison with conventional & reported materials ----
heading(doc, "4.7  Comparison with Conventional and Reported Materials", level=2)
P(doc, "To judge how much the results matter in practice, Table 12 sets the "
  "best-performing formulation (S3) and the binder-less fine panel (S2) against "
  "conventional insulation and against water-hyacinth composites from the "
  "literature. The conductivities of S2 (0.0577 W/m\u00b7K) and S3 (0.0608 "
  "W/m\u00b7K) are higher than those of optimised commercial insulants such as "
  "glass wool (\u2248 0.034 W/m\u00b7K [7]) and EPS (\u2248 0.035\u20130.040 "
  "W/m\u00b7K), but they sit in the same insulation-grade band and match or beat "
  "several reported water-hyacinth composites \u2014 for instance the WH\u2013"
  "cement composite of Philip and Rakendu (0.0765 W/m\u00b7K [3]) and the WH\u2013"
  "bagasse\u2013epoxy composite of Anjani et al. (0.1987 W/m\u00b7K [6]). What "
  "really sets the present panels apart is that they reach this performance with "
  "nothing more than a biodegradable corn-starch binder \u2014 no synthetic resin, "
  "no cement \u2014 which gives them a clear edge on biodegradability and embodied "
  "carbon.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Comparison of the present bio-composites with conventional and "
        "reported insulation materials.")
make_table(doc,
    ["Material", "K (W/m\u00b7K)", "Binder / Matrix", "Reference"],
    [["Glass wool", "\u2248 0.034", "Inorganic (synthetic)", "Jeon et al. [7]"],
     ["EPS foam", "0.035\u20130.040", "Polystyrene (synthetic)", "Typical [7]"],
     ["PU foam", "0.022\u20130.028", "Polyurethane (synthetic)", "Typical"],
     ["WH binder-less board", "0.047\u20130.065", "None", "Salas-Ruiz et al. [2]"],
     ["WH\u2013cement composite", "0.0765", "Cement", "Philip & Rakendu [3]"],
     ["WH\u2013bagasse\u2013epoxy", "0.1987", "Epoxy (synthetic)", "Anjani et al. [6]"],
     ["S2 (this work)", "0.0577", "None (water only)", "Present study"],
     ["S3 (this work)", "0.0608", "Corn starch (90:10)", "Present study"],
     ["S4 (this work)", "0.1846", "Corn starch (70:30)", "Present study"]],
    col_widths=[1.9, 1.2, 1.9, 1.6], font_size=9.5,
    highlight={(6,0):'E2EFDA',(6,1):'E2EFDA',(7,0):'E2EFDA',(7,1):'E2EFDA'})
P(doc, "The comparison confirms that the 90:10 corn-starch panel (S3) is a "
  "credible, fully bio-based insulation candidate: it stays within the "
  "insulation-grade band of binder-less water-hyacinth boards while being far "
  "tougher mechanically, and it beats both cement- and epoxy-bound water-hyacinth "
  "composites on thermal conductivity.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ---- 4.8 SEM Morphological Analysis (with actual micrographs) ----
heading(doc, "4.8  SEM Morphological Analysis", level=2)
P(doc, "Scanning electron microscopy was carried out on the fine binder-less panel "
  "(S1) and the 90:10 starch-bound panel (S3) to draw a direct line between their "
  "microstructure and the macroscopic thermal, mechanical and moisture behaviour "
  "measured in Sections 4.1\u20134.5. Several micrographs were taken at different "
  "magnifications for each sample; the full sets are shown and discussed below.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "4.8.1  Sample S1 \u2014 Fine Binder-less Panel", level=3)
P(doc, "The SEM micrographs of S1 (Figures 20\u201324) show a very open, porous "
  "fibre network, just what you would expect from a binder-less lignocellulosic "
  "composite. At low magnification, randomly oriented biomass fibres and fragments "
  "sit amid large inter-particle voids, with no continuous matrix phase in sight "
  "\u2014 clear proof that there is no binder. The fibres keep their naturally "
  "tubular, sponge-like cross-section, a signature of water-hyacinth petiole "
  "tissue, which adds yet more trapped air both inside the fibre lumen and between "
  "the fibres. Zoom in and the fibre surfaces look rough and uncoated, with cell-"
  "wall microfibrils on show; the fibres touch only by friction and interlocking, "
  "which fits the fairly modest compressive strength of 94.4 kPa measured for S1. "
  "That generous void space is exactly what gives S1 its excellent insulating "
  "performance (K = 0.0577 W/m\u00b7K), as still air trapped in the pores holds "
  "back heat conduction. The same open structure is also why it drinks up so much "
  "water (522.86 %): the interconnected pores let water in freely.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img_row(doc,
    [(f'{PHOTOS}/sem_s1_00.jpg', '(a) S1 \u2014 overview (low magnification)'),
     (f'{PHOTOS}/sem_s1_01.jpg', '(b) S1 \u2014 fibre network detail')],
    total_width_in=6.2,
    cap_text="SEM micrographs of S1 (fine binder-less panel): (a) low-magnification "
             "overview showing open porous fibre network; (b) detail highlighting "
             "large inter-particle voids and randomly oriented fibres.")

img_row(doc,
    [(f'{PHOTOS}/sem_s1_02.jpg', '(c) S1 \u2014 fibre surface morphology'),
     (f'{PHOTOS}/sem_s1_03.jpg', '(d) S1 \u2014 cellular fibre cross-section')],
    total_width_in=6.2,
    cap_text="SEM micrographs of S1 (continued): (c) uncoated fibre surface with "
             "exposed cell-wall microfibrils; (d) natural tubular cellular "
             "cross-section of water-hyacinth fibre adding trapped-air lumen volume.")

img(doc, f'{PHOTOS}/sem_s1_04.jpg', 5.0,
    "SEM micrograph of S1 (high magnification) \u2014 inter-fibre contact zone "
    "showing purely mechanical interlocking with no binder bridging between fibres.")

heading(doc, "4.8.2  Sample S3 \u2014 Fine 90:10 Corn-Starch Panel", level=3)
P(doc, "The SEM micrographs of S3 (Figures 25\u201330) tell a very different story. "
  "At low magnification a thin but continuous gelatinised-starch matrix can be "
  "seen coating the fibre surfaces and bridging the points where fibres meet, "
  "which confirms that the corn-starch binder gelatinised and formed a film as "
  "intended. The fibres are still the dominant structural phase and a good deal "
  "of inter-fibre porosity survives, which is why S3 still manages an "
  "insulation-grade thermal conductivity (K = 0.0608 W/m\u00b7K), only a touch "
  "above the binder-less S1 (0.0577 W/m\u00b7K). At higher magnification the "
  "starch film shows up as discrete bridges and menisci at fibre\u2013fibre "
  "junctions, acting as load-transfer points that explain the much higher "
  "compressive strength of 186.0 kPa, a 97 % gain over S1.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "The starch bridges also partly seal the open pore throats, trimming the "
  "accessible pore volume and bringing water absorption down a little, from "
  "522.86 % (S1) to 507.14 % (S3). The starch film is itself hygroscopic, which is "
  "why S3 holds a higher equilibrium moisture content (23.93 % against 18.67 %), "
  "in keeping with the literature on starch\u2013fibre composites [5], [8]. There "
  "are no fibre pull-out cavities to be seen, which suggests the gelatinised "
  "starch sticks well to the hydroxyl-rich cellulosic fibre surface. All told, the "
  "SEM analysis shows that 10 wt% corn starch lays down a thin, well-spread binder "
  "film that strengthens the inter-fibre bonds without really clogging the "
  "macropores that do the insulating \u2014 which is what makes the 90:10 "
  "formulation perform so well on both fronts.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

img_row(doc,
    [(f'{PHOTOS}/sem_s2_05.jpg', '(a) S3 \u2014 overview (low magnification)'),
     (f'{PHOTOS}/sem_s2_06.jpg', '(b) S3 \u2014 starch-coated fibre network')],
    total_width_in=6.2,
    cap_text="SEM micrographs of S3 (90:10 corn-starch panel): (a) low-magnification "
             "overview showing starch-infiltrated porous network; (b) gelatinised "
             "starch film coating fibre surfaces.")

img_row(doc,
    [(f'{PHOTOS}/sem_s2_07.jpg', '(c) S3 \u2014 starch binder bridges'),
     (f'{PHOTOS}/sem_s2_08.jpg', '(d) S3 \u2014 fibre\u2013matrix interface')],
    total_width_in=6.2,
    cap_text="SEM micrographs of S3 (continued): (c) starch menisci bridging "
             "inter-fibre junctions as load-transfer points; (d) fibre\u2013starch "
             "interface showing good adhesion with no pull-out cavities.")

img_row(doc,
    [(f'{PHOTOS}/sem_s2_09.jpg', '(e) S3 \u2014 retained macro-porosity'),
     (f'{PHOTOS}/sem_s2_10.jpg', '(f) S3 \u2014 high-magnification binder film')],
    total_width_in=6.2,
    cap_text="SEM micrographs of S3 (continued): (e) macro-pores retained despite "
             "10 wt% starch addition, consistent with insulation-grade K; "
             "(f) high-magnification view of the thin continuous starch film "
             "bridging adjacent fibres.")

heading(doc, "4.8.3  Comparative Microstructural Observations", level=3)
P(doc, "Table 13 pulls together the main microstructural observations from the SEM "
  "work and lines them up against the measured macroscopic properties of S1 and "
  "S3.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
tbl_cap(doc, "Summary of SEM microstructural observations and correlation with "
        "macroscopic properties (S1 vs S3).")
make_table(doc,
    ["Feature", "S1 \u2014 Fine Binder-less", "S3 \u2014 90:10 Starch",
     "Property Implication"],
    [["Binder / matrix",
      "Absent; purely mechanical inter-fibre contact",
      "Thin continuous starch film; discrete bridges at junctions",
      "S3 compressive strength 97 % higher"],
     ["Pore structure",
      "Large, interconnected open pores; very porous",
      "Pores partially bridged; macro-porosity retained",
      "Both insulation-grade; S3 K only 5 % higher than S1"],
     ["Fibre surface",
      "Rough, uncoated; exposed microfibrils visible",
      "Smooth starch coating on most fibre surfaces",
      "Improved fibre\u2013matrix adhesion in S3"],
     ["Fibre lumen",
      "Open tubular lumen; additional trapped-air volume",
      "Lumen partially infiltrated by starch gel",
      "S1 marginally lower K due to more trapped air"],
     ["Inter-fibre bonding",
      "Friction and mechanical interlocking only",
      "Starch adhesive menisci \u2014 load-transfer points",
      "Explains higher UCT of S3 (186 vs 94.4 kPa)"],
     ["Pore throat sealing",
      "All pores fully accessible to water",
      "Some pore throats sealed by starch film",
      "S3 WA slightly lower (507 % vs 523 %)"]],
    col_widths=[1.25, 1.6, 1.7, 2.0], font_size=9.0, zebra=True)
P(doc, "Together, the SEM evidence backs the conclusion that 10 wt% corn starch is "
  "the optimal binder loading for this system: it adds just enough bonding to "
  "roughly double the compressive strength while leaving enough open porosity to "
  "keep the thermal conductivity insulation-grade. Taking the starch up to 30 wt% "
  "(S4) would instead give a much denser, more continuous matrix that chokes the "
  "macro-pores, which fits the three-fold rise in K and the oddly low strength "
  "caused by the plasticised, moisture-laden starch phase seen for S4 in "
  "Section 4.4.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
#  CHAPTER 5 — CONCLUSIONS AND SCOPE FOR FUTURE WORK
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 5", level=1, space_before=0)
heading(doc, "CONCLUSIONS AND SCOPE FOR FUTURE WORK", level=1, space_before=2)

heading(doc, "5.1  Conclusions", level=2)
P(doc, "This project showed that invasive waste aquatic biomass from Dal Lake can "
  "realistically be turned into fully bio-based insulation panels held together "
  "with a natural corn-starch binder. The main conclusions are:",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
for c in [
    "Four bio-composite panels were successfully fabricated from water hyacinth "
    "and water lily; particle size was found to be decisive, as the coarse "
    "binder-less panel (S1) fractured on demoulding while all fine-particle panels "
    "(S2\u2013S4) formed intact specimens.",
    "An optimal binder content of 10 wt% corn starch was identified: the 90:10 "
    "panel (S3) achieved the highest unconfined compressive strength of 186.0 kPa "
    "\u2014 a 97 % improvement over the binder-less fine panel (S2, 94.4 kPa).",
    "Two formulations qualified as insulation-grade per ASTM C168: S2 "
    "(K = 0.0577 W/m\u00b7K) and S3 (K = 0.0608 W/m\u00b7K), both below the "
    "0.065 W/m\u00b7K threshold.",
    "Excess starch was detrimental: the 70:30 panel (S4) showed the lowest "
    "strength (26.8 kPa) and a three-fold higher thermal conductivity "
    "(0.1846 W/m\u00b7K), caused by a weak plasticised matrix and a high moisture "
    "content (41.18 %).",
    "Water absorption was high for all panels (280\u2013658 %) but decreased with "
    "binder fraction, confirming that moisture resistance is the key property "
    "requiring improvement.",
    "Overall, the 90:10 biomass:starch panel (S3) is the best-balanced "
    "formulation, uniquely combining insulation-grade thermal conductivity with "
    "the highest mechanical strength, and is therefore recommended as the optimal "
    "formulation from this study.",
    "SEM analysis confirmed that 10 wt% corn starch forms a thin, well-"
    "distributed binder film that bridges inter-fibre junctions and doubles "
    "compressive strength without appreciably filling the macropores that govern "
    "thermal insulation. The binder-less panel (S1) shows a purely open fibre "
    "network with exposed microfibrils, while S3 retains comparable porosity "
    "with added starch bridges at contact points \u2014 a microstructural "
    "explanation consistent with all the measured property trends.",
]:
    pp = doc.add_paragraph(style='List Number')
    pp.paragraph_format.line_spacing = 1.5
    pp.paragraph_format.space_after = Pt(3)
    r = pp.add_run(c); r.font.name = FONT; r.font.size = Pt(12)

heading(doc, "5.2  Scope for Future Work", level=2)
P(doc, "The following directions are recommended to advance this work towards a "
  "deployable product:", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
for f in [
    "Hydrophobic treatment: apply beeswax, linseed-oil or silicone coatings to "
    "reduce water absorption, following the beeswax approach of Chaireh et al. [5].",
    "Fibre surface modification: investigate mild alkali (NaOH) treatment to "
    "improve fibre\u2013binder adhesion and strength, while avoiding the "
    "micro-void formation noted by Abral et al. [10].",
    "Flame retardancy: incorporate eco-friendly retardants such as borax, boric "
    "acid or ammonium polyphosphate, and evaluate LOI and UL-94 ratings [9].",
    "Complete FTIR characterization to correlate fibre surface chemistry "
    "with the measured mechanical and moisture-resistance properties, "
    "complementing the SEM microstructural analysis completed in this work.",
    "Optimise binder content between 10 % and 30 % to refine the strength\u2013"
    "thermal trade-off, and explore hybrid binders (e.g. starch\u2013chitosan).",
    "Hot-pressing trials to densify panels in a controlled manner and study the "
    "density\u2013conductivity\u2013strength relationship.",
    "End-of-life assessment via pyrolysis for energy recovery, and a comparative "
    "life-cycle assessment (LCA) against EPS, PU foam and glass wool.",
    "Scale-up of the fabrication route (mechanical harvesting, industrial drying, "
    "continuous pressing) toward pilot-scale panel production.",
]:
    pp = doc.add_paragraph(style='List Number')
    pp.paragraph_format.line_spacing = 1.5
    pp.paragraph_format.space_after = Pt(3)
    r = pp.add_run(f); r.font.name = FONT; r.font.size = Pt(12)

heading(doc, "5.3  Scale-Up Strategy", level=2)
P(doc, "The study was done at laboratory scale, but the fabrication route lends "
  "itself naturally to industrial scale-up using established equipment. A "
  "continuous manufacturing line for the panels might run through the following "
  "stages:",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
for s in [
    "Mechanical harvesting: amphibious weed harvesters for large-scale, "
    "continuous collection of aquatic biomass from the lake surface.",
    "Industrial drying: rotary drum dryers to reduce the biomass moisture content "
    "to below 10 % efficiently and at scale.",
    "Refining: hammer mills or disc refiners to process bulk biomass into a "
    "uniform fibrous feedstock of controlled particle size.",
    "Resination: a continuous spray line to apply the gelatinised starch binder "
    "(and, in future, flame-retardant additives) uniformly to the fibre.",
    "Continuous pressing: a multi-daylight press to consolidate and form panels "
    "of standard dimensions at controlled density.",
    "Finishing: automated trimming, sanding and application of a hydrophobic "
    "surface coating, followed by quality inspection and packing.",
]:
    pp = doc.add_paragraph(style='List Number')
    pp.paragraph_format.line_spacing = 1.5
    pp.paragraph_format.space_after = Pt(3)
    r = pp.add_run(s); r.font.name = FONT; r.font.size = Pt(12)
P(doc, "A line like this would turn a continuous waste stream into a saleable "
  "construction product, tying lake-restoration work directly to a workable "
  "manufacturing value chain.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

heading(doc, "5.4  Environmental and Economic Significance", level=2)
P(doc, "The environmental case for this work has two sides. First, putting water "
  "hyacinth and water lily to use gives an invasive weed a productive end \u2014 "
  "left to rot, the weed releases methane, a potent greenhouse gas, and returns "
  "nutrients that keep Dal Lake's eutrophication going. Locking that biomass into "
  "durable panels effectively keeps its carbon out of circulation for as long as "
  "the product lasts. Second, swapping petroleum-derived insulants (EPS, PU foam) "
  "and energy-hungry inorganic ones (glass wool, mineral wool) for a "
  "biodegradable, plant-based panel cuts the embodied energy and carbon of the "
  "insulation considerably.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
P(doc, "The economics are favourable too. The raw material is essentially free and "
  "plentiful, its removal is already paid for under lake-conservation programmes, "
  "and the corn-starch binder is a cheap, food-grade commodity. The processing "
  "route is simple and low-temperature \u2014 drying, grinding, mixing and cold "
  "pressing \u2014 with no exotic chemicals and no high-pressure or "
  "high-temperature equipment, so capital and running costs stay low. At the end "
  "of their life the panels are biodegradable, or they can be pyrolysed for energy "
  "recovery, closing the material loop in the spirit of the circular economy.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
page_break(doc)

print("SEM section + Chapter 5 (with scale-up & significance) added.")



# ─────────────────────────────────────────────────────────────────────────────
#  REFERENCES (IEEE, order of citation)
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "REFERENCES", level=1, space_before=0)
hrule(doc)
references = [
    # [1] — cited in Chapter 1 & 2
    "C. Jaktorn and S. Jiajitsawat, \u201cProduction of thermal insulator from "
    "water hyacinth fiber and natural rubber latex,\u201d J. Ecol. Eng., vol. 22, "
    "no. 7, pp. 134\u2013141, 2021. doi: 10.12911/22998993/138736.",
    # [2]
    "A. Salas-Ruiz, M. del M. Barbero-Barrera, and T. Ruiz-T\u00e9llez, "
    "\u201cMicrostructural and thermo-physical characterization of a water "
    "hyacinth petiole for thermal insulation particle board manufacture,\u201d "
    "Materials, vol. 12, no. 4, p. 560, Feb. 2019. doi: 10.3390/ma12040560.",
    # [3]
    "S. Philip and R. Rakendu, \u201cThermal insulation materials based on water "
    "hyacinth for application in sustainable buildings,\u201d Mater. Today Proc., "
    "vol. 57, pp. 1863\u20131867, 2022. doi: 10.1016/j.matpr.2022.01.062.",
    # [4]
    "Y. Zhou, A. Trabelsi, and M. El Mankibi, \u201cHygrothermal properties of "
    "insulation materials from rice straw and natural binders for buildings,\u201d "
    "Polymers, vol. 14, no. 9, p. 1735, Apr. 2022. doi: 10.3390/polym14091735.",
    # [5]
    "P. Chaireh, N. Meethong, and P. Khomwaen, \u201cNovel composite foam made "
    "from starch and water hyacinth with beeswax coating for food packaging "
    "applications,\u201d Int. J. Biol. Macromol., vol. 165, pp. 1382\u20131391, "
    "Dec. 2020. doi: 10.1016/j.ijbiomac.2020.09.243.",
    # [6]
    "A. Anjani, B. Iskandar, and M. Aziz, \u201cThe utilization of composite "
    "material: water hyacinth and sugarcane bagasse fiber\u2013epoxy for cool box "
    "thermal insulation,\u201d J. Energy Mech. Mater. Manuf. Eng. (JEMMME), "
    "vol. 8, no. 1, pp. 29\u201338, 2023. doi: 10.22219/jemmme.v8i1.25010.",
    # [7]
    "J. Jeon, S. Park, and S. Kim, \u201cA study on insulation characteristics of "
    "glass wool and mineral wool coated with a polysiloxane agent,\u201d Adv. "
    "Mater. Sci. Eng., vol. 2017, art. 3938965, 2017. "
    "doi: 10.1155/2017/3938965.",
    # [8]
    "Syamsuri, A. Wahyudi, and R. Ismail, \u201cSynthesis of water hyacinth/"
    "cassava starch composite as an environmentally friendly plastic "
    "solution,\u201d Equilibrium J. Chem. Eng., vol. 7, no. 2, pp. 57\u201363, "
    "2023. doi: 10.20961/equilibrium.v7i2.70321.",
    # [9]
    "P. Suwanniroj and N. Suppakarn, \u201cWater hyacinth fiber as a bio-based "
    "carbon source for intumescent flame-retardant poly(butylene succinate) "
    "composites,\u201d Polymers, vol. 15, no. 21, p. 4211, Oct. 2023. "
    "doi: 10.3390/polym15214211.",
    # [10]
    "H. Abral, D. Kadriadi, A. Rodawan, S. M. Sapuan, and E. S. Zainudin, "
    "\u201cMechanical properties of water hyacinth fibers\u2013polyester composites "
    "before and after immersion in water,\u201d Mater. Des., vol. 58, "
    "pp. 125\u2013129, Jun. 2014. doi: 10.1016/j.matdes.2014.01.043.",
    # [11]
    "J. Pinto, E. Pereira, A. Tavares, and V. M. Ferreira, \u201cCorn\u2019s cob "
    "as a potential ecological thermal insulation material,\u201d Constr. Build. "
    "Mater., vol. 277, art. 122282, Mar. 2021. "
    "doi: 10.1016/j.conbuildmat.2021.122282.",
    # [12]
    "L. Yang, D. Park, and Z. Qin, \u201cMaterial function of mycelium-based "
    "bio-composite: a review,\u201d J. Fungi, vol. 6, no. 4, p. 282, Nov. 2020. "
    "doi: 10.3390/jof6040282.",
    # [13]
    "A. F. Sahayaraj, M. Muthukrishnan, J. Ramesh, and J. T. W. Jappes, "
    "\u201cFlame retardancy of natural fibre-reinforced polymer composites: "
    "a review of mechanisms, additives, and testing methods,\u201d Polymers, "
    "vol. 15, no. 18, p. 3721, Sep. 2023. doi: 10.3390/polym15183721.",
    # [14]
    "H.-R. Kym\u00e4l\u00e4inen and A.-M. Sj\u00f6berg, \u201cFlax and hemp fibres "
    "as raw materials for thermal insulation,\u201d Build. Environ., vol. 43, "
    "no. 7, pp. 1261\u20131269, Jul. 2008. doi: 10.1016/j.buildenv.2007.03.006.",
    # [15]
    "F. Asdrubali, F. D\u2019Alessandro, and S. Schiavoni, \u201cA review of "
    "unconventional sustainable building insulation materials,\u201d Sustain. "
    "Mater. Technol., vol. 4, pp. 1\u201317, Dec. 2015. "
    "doi: 10.1016/j.susmat.2015.05.002.",
    # [16]
    "T.-T. Chen, J.-T. Ye, Q.-Y. Li, and H.-C. Zhang, \u201cThermal stability "
    "and flame-retardant performance of wood-plastic composites treated with "
    "boron compounds,\u201d Compos. Part B Eng., vol. 193, art. 108010, "
    "Jul. 2020. doi: 10.1016/j.compositesb.2020.108010.",
    # [17]
    "N. A. M. Aridi, S. M. Sapuan, E. S. Zainudin, and A. M. Al-Oqla, "
    "\u201cMechanical and morphological properties of injection-moulded rice "
    "husk polypropylene composites,\u201d Int. J. Polym. Anal. Charact., "
    "vol. 21, no. 4, pp. 305\u2013313, 2016. "
    "doi: 10.1080/1023666X.2016.1148316.",
    # [18]
    "A. Oushabi, S. Sair, Y. Abboud, O. Tanane, and A. El Bouari, "
    "\u201cThermal and mechanical characterization of alkali-treated sugarcane "
    "bagasse-reinforced thermoset composites,\u201d South African J. Chem. Eng., "
    "vol. 40, pp. 104\u2013112, Apr. 2022. doi: 10.1016/j.sajce.2022.02.006.",
    # [19]
    "X. Wang, Z. Li, H. Shi, and Y. Yu, \u201cNatural pineapple-leaf fibre: "
    "a promising material for high-performance composites,\u201d Ind. Crops Prod., "
    "vol. 195, art. 116447, May 2023. doi: 10.1016/j.indcrop.2023.116447.",
    # [20]
    "S. Sair, A. Oushabi, A. Kammouni, O. Tanane, Y. Abboud, and "
    "A. El Bouari, \u201cMechanical and thermal conductivity properties of "
    "hemp fibre reinforced polyurethane composites,\u201d Case Stud. Constr. "
    "Mater., vol. 15, art. e00625, Dec. 2021. doi: 10.1016/j.cscm.2021.e00625.",
    # [21]
    "R. Muthuraj, M. Misra, and A. K. Mohanty, \u201cUnidirectional kenaf "
    "fibre-reinforced reprocessable thermoset biocomposites: effect of fibre "
    "loading and UV/humidity ageing,\u201d Compos. Part A Appl. Sci. Manuf., "
    "vol. 156, art. 106874, May 2022. doi: 10.1016/j.compositesa.2022.106874.",
    # [22]
    "D. O. Oyejobi, J. O. Odeyemi, and C. C. Okonkwo, \u201cCompressive strength "
    "and water absorption of coir fibre-reinforced concrete,\u201d J. Nat. "
    "Fibers, vol. 18, no. 9, pp. 1299\u20131311, 2021 (pub. online 2020). "
    "doi: 10.1080/15440478.2019.1697990.",
    # [23]
    "Q. Zhang, H. Li, and L. Li, \u201cA review of end-of-life management "
    "strategies for natural fibre-reinforced polymer composites,\u201d "
    "J. Clean. Prod., vol. 443, art. 141024, Feb. 2024. "
    "doi: 10.1016/j.jclepro.2024.141024.",
    # [24]
    "H. Binici, O. Aksogan, and C. Demirhan, \u201cMechanical, thermal and "
    "acoustical characterizations of an insulation composite made of bio-based "
    "chicken feather and cotton waste fibres,\u201d Sustain. Mater. Technol., "
    "vol. 35, art. e00549, Jan. 2023. doi: 10.1016/j.susmat.2022.e00549.",
    # [25]
    "H. Li, J. Zhang, and Y. Liu, \u201cHygrothermal ageing behaviour and "
    "moisture buffering capacity of bamboo-fibre-reinforced composites,\u201d "
    "Constr. Build. Mater., vol. 271, art. 121555, Feb. 2021. "
    "doi: 10.1016/j.conbuildmat.2020.121555.",
    # [26]
    "J. Liu, L. Chen, and Y. Zhao, \u201cRecent advances and challenges in "
    "cellulose aerogel-based thermal insulation materials: a review,\u201d "
    "Carbohydr. Polym., vol. 324, art. 121533, Jan. 2024. "
    "doi: 10.1016/j.carbpol.2023.121533.",
    # [27]
    "A. Trabelsi, M. El Mankibi, and F. Michel, \u201cHygrothermal performance "
    "of hemp-lime concrete wall: experimental and numerical study,\u201d "
    "Energy Build., vol. 223, art. 110134, Sep. 2020. "
    "doi: 10.1016/j.enbuild.2020.110134.",
    # [28]
    "K. Pawlowski, A. Strzalkowska, and B. Chojnacka, \u201cExploring advancements "
    "in bio-based composites for thermal insulation: a systematic review,\u201d "
    "Sustainability, vol. 17, no. 3, p. 1143, Jan. 2025. "
    "doi: 10.3390/su17031143.",
    # [29]
    "L. Cosentino, P. Fernandes, and R. Mateus, \u201cA review of natural bio-based "
    "insulation materials,\u201d Energies, vol. 16, no. 13, p. 4926, Jun. 2023. "
    "doi: 10.3390/en16134926.",
    # ASTM standards
    "ASTM C168, Standard Terminology Relating to Thermal Insulation, ASTM "
    "International, West Conshohocken, PA, USA.",
    "ASTM D4442, Standard Test Methods for Direct Moisture Content Measurement of "
    "Wood and Wood-Based Materials, ASTM International.",
    "ASTM D1037, Standard Test Methods for Evaluating Properties of Wood-Base "
    "Fiber and Particle Panel Materials, ASTM International.",
    "ASTM D570, Standard Test Method for Water Absorption of Plastics, ASTM "
    "International.",
    "ASTM D5334, Standard Test Method for Determination of Thermal Conductivity of "
    "Soil and Soft Rock by Thermal Needle Probe Procedure, ASTM International.",
]
for i, ref in enumerate(references, 1):
    pp = doc.add_paragraph()
    pp.paragraph_format.line_spacing = 1.5
    pp.paragraph_format.space_after = Pt(6)
    pp.paragraph_format.left_indent = Inches(0.5)
    pp.paragraph_format.first_line_indent = Inches(-0.5)
    r1 = pp.add_run(f"[{i}]\t"); r1.font.name = FONT; r1.font.size = Pt(11)
    r2 = pp.add_run(ref); r2.font.name = FONT; r2.font.size = Pt(11)
page_break(doc)

print("References added.")

# ─────────────────────────────────────────────────────────────────────────────
#  ANNEXURE B — CURRICULUM VITAE (all 3 students)
# ─────────────────────────────────────────────────────────────────────────────
def cv(doc, name, enrol):
    P(doc, "CURRICULUM VITAE", size=14, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
    hrule(doc)
    # Two-column table: info on left, passport photo box on right
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # remove borders
    tblPr2 = tbl._tbl.tblPr
    borders2 = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e2 = OxmlElement(f'w:{edge}'); e2.set(qn('w:val'),'none'); borders2.append(e2)
    tblPr2.append(borders2)
    info_cell = tbl.rows[0].cells[0]
    info_cell.width = Inches(4.2)
    photo_cell = tbl.rows[0].cells[1]
    photo_cell.width = Inches(2.0)
    rows = [
        ("NAME", name),
        ("ENROLMENT NO.", enrol),
        ("DATE OF BIRTH", "______________________"),
        ("FATHER\u2019S NAME", "______________________"),
        ("PERMANENT ADDRESS", "______________________"),
        ("PHONE NUMBER", "______________________"),
        ("EMAIL ADDRESS", "______________________"),
    ]
    t = info_cell.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        c = t.add_row().cells
        c[0].width = Inches(2.0); c[1].width = Inches(2.2)
        p0 = c[0].paragraphs[0]; r0 = p0.add_run(k)
        r0.font.name = FONT; r0.font.size = Pt(12); r0.bold = True
        p0.paragraph_format.space_after = Pt(4)
        p1 = c[1].paragraphs[0]; r1 = p1.add_run(": " + v)
        r1.font.name = FONT; r1.font.size = Pt(12)
        p1.paragraph_format.space_after = Pt(4)
    # remove inner table borders
    tblPr3 = t._tbl.tblPr
    borders3 = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e3 = OxmlElement(f'w:{edge}'); e3.set(qn('w:val'),'none'); borders3.append(e3)
    tblPr3.append(borders3)
    # Passport photo box in right cell
    pp_para = photo_cell.paragraphs[0]
    pp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp_para.paragraph_format.space_before = Pt(4)
    pp_para.paragraph_format.space_after = Pt(4)
    pPr_ph = pp_para._p.get_or_add_pPr()
    pbdr_ph = OxmlElement('w:pBdr')
    for side in ('top','bottom','left','right'):
        b_ph = OxmlElement(f'w:{side}')
        b_ph.set(qn('w:val'),'single'); b_ph.set(qn('w:sz'),'12')
        b_ph.set(qn('w:space'),'8'); b_ph.set(qn('w:color'),'444444')
        pbdr_ph.append(b_ph)
    pPr_ph.append(pbdr_ph)
    run_add(pp_para, "\n\n  Passport\n  Photo\n\n", size=11, italic=True, color=GREY)

    P(doc, "QUALIFICATION", size=12, bold=True, space_before=10, space_after=4)
    make_table(doc,
        ["Examination", "Year", "Institute / Board", "Result"],
        [["B.Tech (Chemical Engineering)", "2026", "NIT Srinagar", "Awaited"],
         ["12th (Senior Secondary)", "______", "______________", "____ %"],
         ["10th (Secondary)", "______", "______________", "____ %"]],
        col_widths=[2.4,0.9,2.1,1.0], font_size=11, zebra=False)
    P(doc, "DECLARATION: I hereby declare that the information furnished above is "
      "true to the best of my knowledge.",
      size=11, italic=True, space_before=12, space_after=18,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_add(pp, "Signature: ____________________", size=12)
    pp2 = doc.add_paragraph(); pp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_add(pp2, name, size=12, bold=True)

P(doc, "BRIEF BIO DATA OF THE CANDIDATES", size=14, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=2)
P(doc, "(Annexure B)", size=11, italic=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
cv(doc, "Suchetan Karloopia", "2022BCHE019")
page_break(doc)
cv(doc, "Miran Haider", "2022BCHE027")
page_break(doc)
cv(doc, "Akshita Sen", "2022BCHE037")
page_break(doc)

print("Brief Bio Data (CVs) added.")

# ─────────────────────────────────────────────────────────────────────────────
#  APPENDICES
# ─────────────────────────────────────────────────────────────────────────────
P(doc, "APPENDICES", size=15, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT, space_after=10)
hrule(doc)
P(doc, "This section contains supplementary material that supports but is not part "
  "of the main body of the report.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

P(doc, "APPENDIX A \u2014 PLAGIARISM CERTIFICATE", size=13, bold=True,
  color=ACCENT, space_before=6, space_after=10)
P(doc, "This is to certify that the project report entitled \u201cDevelopment and "
  "Characterization of Bio-Composites from Waste Aquatic Biomass for Sustainable "
  "Insulation\u201d has been checked for plagiarism using approved similarity-"
  "detection software, and the similarity index is within the permissible limit "
  "prescribed by the Institute.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
P(doc, "The plagiarism-check report generated by the software is to be obtained "
  "from the project supervisor and attached here.",
  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)
P(doc, "Similarity Index: __________ %", size=12, bold=True, space_after=18)
# placeholder box (NOT a numbered figure)
_pp = doc.add_paragraph()
_pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_pp.paragraph_format.space_before = Pt(18)
_pp.paragraph_format.space_after = Pt(18)
_pPr = _pp._p.get_or_add_pPr()
_pbdr = OxmlElement('w:pBdr')
for _side in ('top','bottom','left','right'):
    _b = OxmlElement(f'w:{_side}')
    _b.set(qn('w:val'),'single'); _b.set(qn('w:sz'),'6')
    _b.set(qn('w:space'),'8'); _b.set(qn('w:color'),'AAAAAA')
    _pbdr.append(_b)
_pPr.append(_pbdr)
run_add(_pp, "[ Plagiarism similarity report \u2014 to be obtained from the "
        "supervisor and attached here ]", size=11, italic=True, color=GREY)
pp = doc.add_paragraph(); pp.paragraph_format.space_before = Pt(24)
run_add(pp, "Signature of Supervisor: ____________________", size=12)
P(doc, "Dr. Fasil Qayoom Mir", size=12, bold=True, space_before=8, space_after=0)
P(doc, "Head & Associate Professor, Department of Chemical Engineering, NIT Srinagar",
  size=11, space_after=0)

print("Appendices (Plagiarism Certificate) added.")



# ═════════════════════════════════════════════════════════════════════════════
#  PAGE NUMBERING  (Roman i, ii, ... for preliminary;  Arabic 1, 2, ... for body)
# ═════════════════════════════════════════════════════════════════════════════
prelim = doc.sections[0]
body   = doc.sections[1]

# Preliminary section: lowercase roman, title page un-numbered
prelim.different_first_page_header_footer = True
set_number_format(prelim, 'lowerRoman', start=1)
f = prelim.footer
f.is_linked_to_previous = False
fp = f.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.line_spacing = 1.0
r = add_page_number_field(fp)
r.font.name = FONT; r.font.size = Pt(11)
# leave first-page (title) footer empty
_ = prelim.first_page_footer.paragraphs[0]

# Body section: arabic starting at 1
set_number_format(body, 'decimal', start=1)
bf = body.footer
bf.is_linked_to_previous = False
bfp = bf.paragraphs[0]
bfp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bfp.paragraph_format.line_spacing = 1.0
rb = add_page_number_field(bfp)
rb.font.name = FONT; rb.font.size = Pt(11)

# ═════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═════════════════════════════════════════════════════════════════════════════
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"\nReport saved to: {OUTPUT}")
print(f"Total tables: {len(doc.tables)}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
