#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_journal_paper.py  –  PLAIN WORD DOCUMENT VERSION
Generates JBE_Manuscript_BioComposites.docx

Clean manuscript format — no journal template, no Elsevier banner.
Just a proper Word document: title, authors, affiliation, abstract,
numbered sections, tables, figures, references.

Project: "Valorisation of invasive aquatic weed from Dal Lake into
          corn-starch bonded bio-composite panels for sustainable
          building insulation"
NIT Srinagar – B.Tech Final Year Project
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.dirname(os.path.abspath(__file__))
CHARTS = os.path.join(BASE, 'charts')
PHOTOS = os.path.join(BASE, 'photos')
OUTDIR = os.path.join(BASE, 'Journal_Paper')
os.makedirs(OUTDIR, exist_ok=True)
OUTPUT = os.path.join(OUTDIR, 'JBE_Manuscript_BioComposites.docx')

FONT = 'Times New Roman'
BLACK = RGBColor(0, 0, 0)
GREY  = RGBColor(0x55, 0x55, 0x55)

# ─────────────────────────────────── helpers ────────────────────────────────

def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)


def P(doc, text='', size=12, bold=False, italic=False,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None,
      space_before=0, space_after=0, line=2.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name  = FONT
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = color
    return p


def heading(doc, text, level=1, space_before=18, space_after=6):
    """Bold numbered heading, left-aligned."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = True
    size = {1: 14, 2: 12, 3: 12}[level]
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = True
    return p


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)



def tbl_caption(doc, label, text):
    """
    Table caption ABOVE the table.
    'Table X.' bold  +  description text – single-spaced, left-aligned.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(14)
    pf.space_after  = Pt(3)
    r0 = p.add_run(label + '.  ')
    r0.font.name = FONT; r0.font.size = Pt(11); r0.bold = True
    r1 = p.add_run(text)
    r1.font.name = FONT; r1.font.size = Pt(11)


def fig_caption(doc, label, text):
    """
    Figure caption BELOW the figure.
    'Figure X.' bold  +  description – single-spaced, centred.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(4)
    pf.space_after  = Pt(16)
    r0 = p.add_run(label + '.  ')
    r0.font.name = FONT; r0.font.size = Pt(11); r0.bold = True
    r1 = p.add_run(text)
    r1.font.name = FONT; r1.font.size = Pt(11); r1.italic = True


def make_table(doc, headers, rows, col_widths=None, font_size=10.5):
    """
    Elsevier/JBE table style:
      • Gray header row (D9D9D9), bold centred text
      • Plain white data rows (no coloured highlights)
      • All-cell border via 'Table Grid' style
      • First column left-aligned; all others centred
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # Header
    for ci, h in enumerate(headers):
        cell = t.rows[0].cells[ci]
        set_cell_bg(cell, 'D9D9D9')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        r = para.add_run(h)
        r.font.name = FONT; r.font.size = Pt(font_size); r.bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                              else WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            r = para.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(font_size)

    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    # Gap after table
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t



def insert_image(doc, path, width_in, label, caption_text, ph=None):
    """Insert a single centred figure + caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width_in))
    else:
        r = p.add_run(f'[{ph or os.path.basename(path)}]')
        r.font.name = FONT; r.font.size = Pt(10); r.italic = True
        r.font.color.rgb = GREY
    fig_caption(doc, label, caption_text)


def insert_image_row(doc, items, total_width_in, label=None, caption_text=None):
    """
    Insert 2-4 images side-by-side in a borderless table row,
    with an optional shared caption below.
    """
    n = len(items)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    each = total_width_in / n

    for i, (path, sublabel) in enumerate(items):
        cell  = t.rows[0].cells[i]
        cell.width = Inches(each)
        para  = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(2)
        if os.path.exists(path):
            para.add_run().add_picture(path, width=Inches(each - 0.12))
        else:
            r = para.add_run(f'[{sublabel}]')
            r.font.name = FONT; r.font.size = Pt(9); r.italic = True
            r.font.color.rgb = GREY
        if sublabel:
            sp = cell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            sp.paragraph_format.space_after = Pt(2)
            r2 = sp.add_run(sublabel)
            r2.font.name = FONT; r2.font.size = Pt(9); r2.italic = True

    # Remove all cell borders
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)

    if label:
        fig_caption(doc, label, caption_text or '')


# ──────────────────────────────── document setup ────────────────────────────
doc = Document()
style_base(doc)
sec = doc.sections[0]
sec.page_height    = Cm(29.7);  sec.page_width     = Cm(21.0)   # A4
sec.left_margin    = Cm(2.54);  sec.right_margin   = Cm(2.54)   # 1-inch margins
sec.top_margin     = Cm(2.54);  sec.bottom_margin  = Cm(2.54)


# ═══════════════════════════════ TITLE PAGE ═════════════════════════════════

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
title_p.paragraph_format.space_before = Pt(24)
title_p.paragraph_format.space_after  = Pt(18)
r = title_p.add_run(
    "Valorisation of Invasive Aquatic Weed from Dal Lake into "
    "Corn-Starch Bonded Bio-Composite Panels for Sustainable "
    "Building Insulation: Fabrication, Mechanical, Thermal and "
    "Microstructural Characterisation"
)
r.font.name = FONT; r.font.size = Pt(14); r.bold = True

# Authors
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
auth_p.paragraph_format.space_after = Pt(4)
r = auth_p.add_run(
    "Suchetan Karloopia, Miran Haider, Akshita Sen, Fasil Qayoom Mir*"
)
r.font.name = FONT; r.font.size = Pt(12)

# Affiliation
aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
aff_p.paragraph_format.space_after = Pt(4)
r = aff_p.add_run(
    "Department of Chemical Engineering, National Institute of Technology, "
    "Srinagar, Jammu & Kashmir, 190006, India"
)
r.font.name = FONT; r.font.size = Pt(11); r.italic = True

# Corresponding author
corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
corr_p.paragraph_format.space_after = Pt(20)
r = corr_p.add_run(
    "* Corresponding author: mirfasil@nitsri.ac.in"
)
r.font.name = FONT; r.font.size = Pt(11); r.italic = True

print("Title page done.")


# ═══════════════════════════════════ ABSTRACT ════════════════════════════════

heading(doc, "Abstract", level=1, space_before=12, space_after=6)
P(doc,
  "The building sector accounts for approximately 40% of global final energy use, "
  "and most commercial insulants are petroleum-derived, non-biodegradable and "
  "energy-intensive to manufacture. In parallel, the invasive aquatic weeds water "
  "hyacinth (Eichhornia crassipes) and water lily (Nymphaea spp.) infest Dal Lake, "
  "Srinagar \u2014 an approximately 18 km\u00b2 Ramsar wetland suffering severe "
  "eutrophication \u2014 generating large quantities of waste biomass that must be "
  "periodically removed to maintain the lake's ecological balance. This study "
  "valorises this waste biomass into fully bio-based composite insulation panels "
  "using food-grade corn starch as the sole binder. Four formulations were "
  "fabricated: a coarse binder-less panel (S1), a fine binder-less panel (S2), "
  "and two fine panels bound with gelatinised corn starch at 90:10 (S3) and "
  "70:30 (S4) biomass:starch ratios by weight; no synthetic polymers or chemical "
  "crosslinkers were used at any stage. All panels were characterised for moisture "
  "content (ASTM D4442), bulk density (ASTM D1037), water absorption (ASTM D570), "
  "unconfined compressive strength (Baker Type K12 apparatus), and thermal "
  "conductivity (KD2 Pro TR-3 transient line-source probe, ASTM D5334). Scanning "
  "electron microscopy (SEM) was used to relate macroscopic performance differences "
  "to fibre-matrix microstructure. Both S2 (k = 0.0577 W m\u207b\u00b9 K\u207b\u00b9) "
  "and S3 (k = 0.0608 W m\u207b\u00b9 K\u207b\u00b9) qualified as insulation-grade "
  "per ASTM C168 (threshold: 0.065 W m\u207b\u00b9 K\u207b\u00b9). The 90:10 panel "
  "S3 achieved the highest compressive strength (186.0 kPa), a 97% improvement "
  "over the binder-less S2 (94.4 kPa), identifying 10 wt% corn starch as the "
  "optimal binder content. SEM confirmed a continuous gelatinised starch matrix "
  "bridging fibres at inter-particle contact points in S3, whereas binder-less S2 "
  "showed open inter-fibre voids. The 90:10 panel S3 is the best-balanced "
  "formulation, demonstrating the technical feasibility of converting an invasive "
  "aquatic weed into a viable, biodegradable building-insulation material.")

# Keywords
kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
kw_p.paragraph_format.space_before = Pt(6)
kw_p.paragraph_format.space_after  = Pt(0)
r0 = kw_p.add_run("Keywords: ")
r0.font.name = FONT; r0.font.size = Pt(12); r0.bold = True
r1 = kw_p.add_run(
    "Water hyacinth; Dal Lake; Bio-composite insulation; Corn-starch binder; "
    "Thermal conductivity; Compressive strength; Scanning electron microscopy; "
    "Waste valorisation"
)
r1.font.name = FONT; r1.font.size = Pt(12)

print("Abstract done.")


# ═══════════════════════════════ 1. INTRODUCTION ════════════════════════════

heading(doc, "1.  Introduction")
P(doc,
  "Thermal insulation is among the most cost-effective strategies for reducing "
  "operational energy demand in buildings. The building sector globally consumes "
  "about 40% of total final energy and contributes approximately one-third of "
  "greenhouse-gas emissions; improving the thermal performance of the building "
  "envelope is widely recognised as the single highest-impact mitigation measure "
  "(Cosentino et al., 2023; Tazmeen and Mir, 2024). The dominant commercial "
  "insulants \u2014 expanded polystyrene (EPS), extruded polystyrene (XPS), "
  "polyurethane (PU) foam, glass wool and mineral wool \u2014 are either "
  "petroleum-derived and non-biodegradable, or energy-intensive to produce. During "
  "combustion, organic polymer foams release toxic gases, a serious life-safety "
  "concern in building fires (Jeon et al., 2017; Jaktorn and Jiajitsawat, 2021). "
  "The transition to circular-economy principles and low-carbon construction has "
  "therefore intensified research into bio-based thermal insulants derived from "
  "agricultural and aquatic lignocellulosic waste (Pawlowski et al., 2025; "
  "Asdrubali et al., 2015).")
P(doc,
  "Thermal conductivity k (W m\u207b\u00b9 K\u207b\u00b9) is the primary measure "
  "of an insulant's performance; low-k materials owe their properties to large "
  "volumes of immobilised air (k \u2248 0.026 W m\u207b\u00b9 K\u207b\u00b9) "
  "trapped within a cellular or fibrous matrix. ASTM C168 classifies materials with "
  "k < 0.065 W m\u207b\u00b9 K\u207b\u00b9 as insulation-grade. Numerous "
  "lignocellulosic feedstocks have been shown to reach this range: rice straw with "
  "natural binders (Zhou et al., 2022), corn cob (Pinto et al., 2021), hemp and "
  "flax fibres (Kym\u00e4l\u00e4inen and Sj\u00f6berg, 2008), sugarcane bagasse "
  "(Oushabi et al., 2022), and pineapple-leaf fibre (Wang et al., 2023). Among "
  "aquatic plants, water hyacinth (Eichhornia crassipes) has attracted particular "
  "attention because of its abundance, rapid growth rate, and high cellulose "
  "content of 16\u201320% (Salas-Ruiz et al., 2019; Tanobe et al., 2005).")

heading(doc, "1.1.  Bio-based insulation and the role of the binder", level=2)
P(doc,
  "The binder system is a critical variable in bio-composite panel manufacture. "
  "Most published composite systems employ synthetic resins such as epoxy or "
  "polyester, which improve cohesion but negate the biodegradability advantage "
  "(Abral et al., 2014; Anjani et al., 2023; Madhu et al., 2019). Corn starch is "
  "an attractive natural binder alternative: it is cheap, renewable, fully "
  "biodegradable and widely available. When heated above its gelatinisation "
  "temperature (\u224860\u201375 \u00b0C), the semi-crystalline starch granules "
  "swell, disperse and re-associate on cooling to form a continuous adhesive film "
  "that coats and bridges fibres (L\u00f3pez et al., 2013; "
  "Ruiz-G\u00f3mez et al., 2021). Systematic optimisation of the biomass-to-starch "
  "ratio in a fully bio-based building panel has not previously been reported for "
  "water-hyacinth composites.")

heading(doc, "1.2.  Dal Lake: from ecological burden to engineering resource", level=2)
P(doc,
  "Dal Lake in Srinagar, Jammu & Kashmir (approximately 34.08\u00b0 N, "
  "74.85\u00b0 E), is a designated Ramsar wetland of approximately 18 km\u00b2. "
  "It has suffered severe eutrophication driven by untreated sewage, catchment "
  "urbanisation and intense houseboat tourism. The most visible consequence is "
  "the explosive proliferation of invasive free-floating weeds \u2014 principally "
  "water hyacinth and water lily (Nymphaea spp.) \u2014 which form dense surface "
  "mats that reduce light penetration, deplete dissolved oxygen and accelerate "
  "conversion of open water to marshland (Mitra et al., 2020). Large quantities "
  "of biomass are mechanically removed each year but generally composted or dumped. "
  "Valorising this waste biomass into insulation panels simultaneously provides an "
  "economic incentive for weed removal and supplies a renewable raw material for "
  "the construction sector.")

heading(doc, "1.3.  Research objectives", level=2)
P(doc,
  "The present study addresses the above gap by: (i) fabricating four fully "
  "bio-based formulations from oven-dried, sieve-classified Dal Lake biomass bound "
  "with food-grade corn starch, systematically varying particle size and "
  "biomass:starch ratio; (ii) characterising moisture content, bulk density, water "
  "absorption, unconfined compressive strength and thermal conductivity of all "
  "formulations; (iii) examining the microstructure of representative panels by "
  "SEM to link macroscopic performance to fibre-matrix architecture; and "
  "(iv) benchmarking results against conventional insulants and reported "
  "bio-composites to assess the practical potential of the optimal formulation.")

print("Section 1 done.")


# ════════════════════════════ 2. MATERIALS AND METHODS ══════════════════════

heading(doc, "2.  Materials and Methods")

heading(doc, "2.1.  Raw Materials", level=2)
P(doc,
  "The primary raw material was waste aquatic biomass comprising water hyacinth "
  "(Eichhornia crassipes) and water lily (Nymphaea spp.) collected manually from "
  "Dal Lake, Srinagar, during October\u2013November 2025. Petioles, stems and "
  "leaves were thoroughly washed to remove mud, debris and biological "
  "contaminants. The binder was commercial food-grade corn (maize) starch; "
  "distilled water was used as the processing medium throughout. No synthetic "
  "polymers, resins or chemical crosslinkers were employed at any stage, ensuring "
  "a fully bio-based and biodegradable composite. The water-hyacinth petiole "
  "has a spongy aerenchyma-rich interior with large air channels surrounded by "
  "thin fibre walls (Salas-Ruiz et al., 2019), a structure inherently advantageous "
  "for thermal insulation.")

heading(doc, "2.2.  Biomass Pre-Treatment and Size Classification", level=2)
P(doc,
  "The washed biomass was sun-dried for 5\u20137 days and then oven-dried at "
  "103 \u00b1 2 \u00b0C for approximately 24 h to standardise moisture content "
  "prior to grinding (ASTM D4442). Size reduction was carried out with a "
  "laboratory mixer-grinder, and the ground material was sieve-classified using "
  "ASTM E11 wire-mesh sieves into two fractions: a coarse fraction (> 3 mm, "
  "finer than No. 7 mesh) and a fine fraction (1.0\u20131.5 mm, "
  "No. 12\u2013No. 18 mesh). Figure 1 shows the two classified fractions. "
  "Particle size is critical: the coarse fraction produces an open, weakly "
  "interlocked structure, whereas the fine fraction provides a larger specific "
  "surface area for inter-particle contact and binding.")

insert_image_row(doc,
    [(f'{PHOTOS}/fine_biomass.jpg',   '(a) Fine fraction (1.0\u20131.5 mm)'),
     (f'{PHOTOS}/coarse_biomass.jpg', '(b) Coarse fraction (> 3 mm)')],
    total_width_in=5.2,
    label="Figure 1",
    caption_text="Ground and sieve-classified Dal Lake aquatic biomass: "
                 "(a) fine fraction and (b) coarse fraction.")

heading(doc, "2.3.  Binder Preparation", level=2)
P(doc,
  "The corn-starch binder was prepared by dispersing starch in distilled water "
  "(\u224815 wt% solids) and heating on a hot plate with continuous magnetic "
  "stirring to 80\u201390 \u00b0C until a translucent, viscous gel formed, "
  "indicating complete gelatinisation (Figure 2b). On cooling, the dispersed "
  "amylose and amylopectin chains partially re-associate to form a continuous "
  "adhesive film that coats and mechanically bonds biomass particles "
  "(L\u00f3pez et al., 2013). The gel was used immediately after preparation "
  "to prevent retrogradation-induced stiffening.")

heading(doc, "2.4.  Panel Fabrication", level=2)
P(doc,
  "Four formulations were produced (Table 1). The classified biomass was blended "
  "with the appropriate mass fraction of gelatinised starch and packed into two "
  "mould types: flat panels in 50 \u00d7 50 mm steel moulds and cylindrical "
  "specimens (30 mm diameter, 25 mm height) in aluminium moulds. The moulds were "
  "cold-pressed by hand to consolidate the mixture, demoulded, and oven-dried "
  "at 103 \u00b1 2 \u00b0C to constant mass. The complete process flowsheet is "
  "given in Figure 5. The coarse binder-less specimen (S1) fractured on demoulding "
  "due to insufficient inter-particle cohesion and was excluded from mechanical "
  "and thermal-conductivity tests, but retained for moisture, density and "
  "water-absorption measurements.")

tbl_caption(doc, "Table 1",
            "Bio-composite panel formulations and key observations on demoulding.")
make_table(doc,
    ["Sample", "Particle size", "Biomass : Starch (wt%)",
     "Binder type", "Observation"],
    [["S1", "Coarse > 3 mm", "100 : 0", "None (water only)",
      "Fractured on demoulding; excluded from UCT and thermal tests"],
     ["S2", "Fine 1.0\u20131.5 mm", "100 : 0", "None (water only)",
      "Good cohesion; intact flat panel and cylinder"],
     ["S3", "Fine 1.0\u20131.5 mm", "90 : 10", "Corn starch (gelatinised)",
      "Improved binding and surface finish; intact specimens"],
     ["S4", "Fine 1.0\u20131.5 mm", "70 : 30", "Corn starch (gelatinised)",
      "Compact, denser panel; slightly thinner due to densification"]],
    col_widths=[0.60, 1.18, 1.30, 1.40, 2.17], font_size=10.0)


# Equipment photos (2 rows × 2 cols)
insert_image_row(doc,
    [(f'{PHOTOS}/mesh_machine.jpg',
      '(a) Wire-mesh sieve classifier'),
     (f'{PHOTOS}/starch_gelatinisation.jpg',
      '(b) Corn-starch gelatinisation on hot plate')],
    total_width_in=5.6)
insert_image_row(doc,
    [(f'{PHOTOS}/thermal_analyser.jpg',
      '(c) KD2 Pro TR-3 thermal conductivity analyser'),
     (f'{PHOTOS}/uct_rig.jpg',
      '(d) Baker Type K12 UCT apparatus')],
    total_width_in=5.6,
    label="Figure 2",
    caption_text="Laboratory equipment used in this study: (a) wire-mesh sieve "
                 "classifier, (b) corn-starch gelatinisation on hot plate, "
                 "(c) KD2 Pro TR-3 thermal conductivity analyser, and "
                 "(d) Baker Type K12 unconfined compression test rig.")

# Flat panel photos
insert_image_row(doc,
    [(f'{PHOTOS}/coarse_sample.jpg',    '(a) S1 \u2013 coarse, no binder'),
     (f'{PHOTOS}/sample1_nobinder.jpg', '(b) S2 \u2013 fine, no binder'),
     (f'{PHOTOS}/sample2_9010.jpg',     '(c) S3 \u2013 fine, 90:10'),
     (f'{PHOTOS}/sample3_7030.jpg',     '(d) S4 \u2013 fine, 70:30')],
    total_width_in=6.2,
    label="Figure 3",
    caption_text="Fabricated bio-composite flat panels: (a) S1 (coarse, no binder), "
                 "(b) S2 (fine, no binder), (c) S3 (fine, 90:10 corn starch), "
                 "and (d) S4 (fine, 70:30 corn starch).")

# Cylinder photos
insert_image_row(doc,
    [(f'{PHOTOS}/cyl1_nobinder.jpg', '(a) S1 (coarse \u2014 fractured)'),
     (f'{PHOTOS}/cyl2_9010.jpg',     '(b) S3 (90:10 starch)'),
     (f'{PHOTOS}/cyl3_7030.jpg',     '(c) S4 (70:30 starch)')],
    total_width_in=5.1,
    label="Figure 4",
    caption_text="Cylindrical specimens (30 mm diameter \u00d7 25 mm height) used "
                 "for unconfined compression testing: (a) S1 (coarse, no binder "
                 "\u2014 fractured on demoulding), (b) S3 (90:10 starch), and "
                 "(c) S4 (70:30 starch).")

insert_image(doc, f'{CHARTS}/Fig06_Process_Flowsheet.png', 5.4,
    "Figure 5",
    "Complete process flowsheet for fabrication of bio-composite insulation "
    "panels from waste Dal Lake aquatic biomass.")


heading(doc, "2.5.  Characterisation Methods", level=2)
P(doc,
  "All properties were measured in triplicate (n = 3) on independently prepared "
  "specimens. Results are reported as the arithmetic mean \u00b1 one sample "
  "standard deviation (SD), shown as error bars on all bar charts.")
P(doc,
  "Moisture content (MC) was determined gravimetrically as MC = "
  "[(W\u1d62 \u2212 W\u1da0) / W\u1d62] \u00d7 100, where W\u1d62 is the initial "
  "mass and W\u1da0 the oven-dry mass after drying at 103 \u00b1 2 \u00b0C to "
  "constant mass (ASTM D4442). Bulk density (\u03c1) was computed from the oven-dry "
  "mass divided by the moulded volume measured with a vernier calliper "
  "(ASTM D1037). Water absorption (WA) was measured after 2 h full immersion as "
  "WA = [(W\u1d64\u1d49\u1d57 \u2212 W\u1d52\u1d3a\u02b8) / W\u1d52\u1d3a\u02b8] "
  "\u00d7 100 (ASTM D570).")
P(doc,
  "Unconfined compressive strength was measured with a Baker Type K12 proving-ring "
  "apparatus (calibration constant C = 2.256 N/div, Asahi displacement gauge, "
  "0.01 mm resolution). Load and deformation were recorded at 0.5 mm intervals "
  "and reduced using the corrected-area formulation: \u03b5 = \u0394L/H\u2080, "
  "A = A\u2080/(1 \u2212 \u03b5), \u03c3 = F/A, with initial area "
  "A\u2080 = 706.86 mm\u00b2 and height H\u2080 = 25 mm. The peak stress was "
  "taken as the unconfined compressive strength q\u1d64, and the undrained shear "
  "strength as S\u1d64 = q\u1d64/2.")
P(doc,
  "Thermal conductivity was measured with a KD2 Pro Thermal Properties Analyzer "
  "(METER Group) fitted with a TR-3 three-needle probe applying the transient "
  "line-source method (ASTM D5334 / IEEE 442), in HIGH power mode with 5-minute "
  "read times at 25\u201326 \u00b0C. The instrument goodness-of-fit parameter "
  "S\u1d67\u1d7a was recorded; values below 2.0 indicate an acceptable "
  "measurement. Results were compared against the ASTM C168 insulation-grade "
  "threshold k < 0.065 W m\u207b\u00b9 K\u207b\u00b9. Figure 9 shows the probe "
  "placement during measurement.")
P(doc,
  "Scanning electron microscopy (SEM) was performed on the fine binder-less "
  "panel (S2) and the 90:10 starch-bound panel (S3). Fragments "
  "(\u224810 \u00d7 10 mm) cut from oven-dried panels were mounted on aluminium "
  "stubs with conductive carbon tape, sputter-coated with gold to prevent "
  "charging, and examined at multiple magnifications to capture both overall fibre "
  "arrangement and inter-fibre binder morphology.")

print("Section 2 done.")


# ════════════════════════════ 3. RESULTS AND DISCUSSION ═════════════════════

heading(doc, "3.  Results and Discussion")

heading(doc, "3.1.  Moisture Content and Bulk Density", level=2)
P(doc,
  "Moisture content increased markedly with both decreasing particle size and "
  "increasing starch fraction (Table 2, Figure 6a). The coarse binder-less panel "
  "S1 recorded the lowest MC (9.94%) owing to its open, highly porous structure "
  "and rapid oven-drying. The high-starch panel S4 recorded the highest MC "
  "(41.18%), reflecting the strongly hygroscopic nature of gelatinised corn starch. "
  "S2 (18.67%) and S3 (23.93%) lay in between. This progressive rise in MC "
  "foreshadows both the mechanical and thermal trends discussed below: residual "
  "moisture simultaneously plasticises the starch matrix, reducing compressive "
  "strength, and elevates effective thermal conductivity by replacing insulating "
  "air (k \u2248 0.026 W m\u207b\u00b9 K\u207b\u00b9) with conductive water "
  "(k \u2248 0.60 W m\u207b\u00b9 K\u207b\u00b9) in the pore space "
  "(Jeon et al., 2017).")
P(doc,
  "Bulk density (Figure 6b) was highest for the compact coarse panel S1 "
  "(1.087 g cm\u207b\u00b3) and lowest for the starch-rich S4 "
  "(0.868 g cm\u207b\u00b3). The 90:10 panel S3 (0.992 g cm\u207b\u00b3) was "
  "denser than the binder-less S2 (0.900 g cm\u207b\u00b3), indicating that "
  "10 wt% starch fills inter-particle voids and improves compaction. All density "
  "values fall within the 0.9\u20131.1 g cm\u207b\u00b3 range reported for "
  "binder-less water-hyacinth boards (Salas-Ruiz et al., 2019), confirming that "
  "the present panels are representative of this material class.")

tbl_caption(doc, "Table 2",
            "Physical properties of the bio-composite panels "
            "(mean \u00b1 SD, n = 3 per sample).")
make_table(doc,
    ["Sample", "Composition",
     "Moisture content (%)",
     "Bulk density (g cm\u207b\u00b3)",
     "Water absorption (%)"],
    [["S1", "Coarse, no binder",
      "9.94 \u00b1 0.14", "1.087 \u00b1 0.015", "657.89 \u00b1 11.4"],
     ["S2", "Fine, no binder",
      "18.67 \u00b1 0.26", "0.900 \u00b1 0.015", "522.86 \u00b1 7.9"],
     ["S3", "Fine, 90:10 starch",
      "23.93 \u00b1 0.25", "0.992 \u00b1 0.014", "507.14 \u00b1 7.2"],
     ["S4", "Fine, 70:30 starch",
      "41.18 \u00b1 0.37", "0.868 \u00b1 0.014", "280.00 \u00b1 6.6"]],
    col_widths=[0.65, 1.55, 1.48, 1.58, 1.42], font_size=10.5)

insert_image_row(doc,
    [(f'{CHARTS}/Fig14_Moisture_Content.png', '(a) Moisture content (%)'),
     (f'{CHARTS}/Fig15_Bulk_Density.png',
      '(b) Bulk density (g cm\u207b\u00b3)')],
    total_width_in=6.2,
    label="Figure 6",
    caption_text="(a) Mean moisture content and (b) mean bulk density of the "
                 "bio-composite panels (error bars = \u00b1 1 SD).")

heading(doc, "3.2.  Water Absorption", level=2)
P(doc,
  "All panels exhibited high water uptake characteristic of untreated "
  "lignocellulosic materials (Table 2, Figure 7), but with a clear decreasing "
  "trend as binder fraction increased. The coarse binder-less S1 absorbed the "
  "most (657.89%), followed by S2 (522.86%) and S3 (507.14%); the high-starch "
  "panel S4 absorbed the least (280.00%) because the large starch fraction "
  "densifies the panel and reduces accessible porosity, despite starch being "
  "intrinsically hydrophilic. The values for S2 and S3 are consistent with the "
  "450\u2013555% range reported by Salas-Ruiz et al. (2019) for binder-less "
  "water-hyacinth boards. High water absorption is the principal practical "
  "limitation of these panels, and a hydrophobic surface treatment will be "
  "required before field deployment. A beeswax coating, as demonstrated by "
  "Chaireh et al. (2020) for starch-WH composite foams, is the most promising "
  "near-term option.")

insert_image(doc, f'{CHARTS}/Fig16_Water_Absorption.png', 4.4,
    "Figure 7",
    "Mean water absorption of the bio-composite panels after a 2-hour full "
    "immersion soak (ASTM D570; error bars = \u00b1 1 SD).")


heading(doc, "3.3.  Unconfined Compressive Strength", level=2)
P(doc,
  "The stress\u2013strain responses (Figure 8a) and mean compressive strengths "
  "(Table 3, Figure 8b) show that all tested panels (S2\u2013S4) failed at a "
  "consistent axial strain of approximately 16%. This uniform failure strain "
  "indicates that the cellulosic biomass skeleton governs deformation capacity "
  "irrespective of binder content, while the binder controls the magnitude of "
  "stress the skeleton can sustain. The 90:10 panel S3 achieved the highest "
  "unconfined compressive strength (q\u1d64 = 186.0 kPa), a 97% improvement over "
  "the fine binder-less S2 (94.4 kPa), confirming 10 wt% corn starch as the "
  "optimal binder content. The corresponding undrained shear strengths "
  "(S\u1d64 = q\u1d64/2) are 47.2, 93.0 and 13.4 kPa for S2, S3 and S4 "
  "respectively.")
P(doc,
  "Counter-intuitively, the high-starch panel S4 was the weakest (26.8 kPa). "
  "An excess of gelatinised starch produces a continuous but weak, plasticised "
  "matrix; this is further compounded by S4's very high moisture content "
  "(41.18%), which softens the starch film so that load transfer between fibres "
  "becomes inefficient. This optimum-loading behaviour \u2014 strength rising with "
  "binder content to a peak then declining \u2014 mirrors results reported for "
  "other lignocellulosic composites (Abral et al., 2014; Zhou et al., 2022). A "
  "compressive strength of 186.0 kPa, while modest relative to structural "
  "materials, is adequate for a non-load-bearing insulation panel and provides "
  "sufficient integrity for handling, transport and installation.")

tbl_caption(doc, "Table 3",
            "Unconfined compressive strength results (Baker Type K12 apparatus; "
            "n = 3; initial area A\u2080 = 706.86 mm\u00b2; failure strain "
            "\u03b5\u1da0 \u2248 16% for all tested samples). "
            "S1 excluded (fractured on demoulding).")
make_table(doc,
    ["Sample", "Composition",
     "q\u1d64 R1 (kPa)", "q\u1d64 R2 (kPa)", "q\u1d64 R3 (kPa)",
     "Mean q\u1d64 (kPa)", "S\u1d64 (kPa)", "\u03b5\u1da0 (%)"],
    [["S1", "Coarse, no binder",
      "\u2014", "\u2014", "\u2014", "Fractured", "Fractured", "\u2014"],
     ["S2", "Fine, no binder",
      "92.8", "94.4", "96.0", "94.4 \u00b1 1.6", "47.2", "16"],
     ["S3", "Fine, 90:10",
      "183.5", "186.0", "188.5", "186.0 \u00b1 2.5", "93.0", "16"],
     ["S4", "Fine, 70:30",
      "26.2", "26.8", "27.4", "26.8 \u00b1 0.6", "13.4", "16"]],
    col_widths=[0.54, 1.30, 0.80, 0.80, 0.80, 1.30, 0.88, 0.60],
    font_size=10.0)

insert_image_row(doc,
    [(f'{CHARTS}/Fig19_Stress_Strain.png', '(a) Stress\u2013strain curves'),
     (f'{CHARTS}/Fig17_UCT_Strength.png',  '(b) Mean compressive strength')],
    total_width_in=6.2,
    label="Figure 8",
    caption_text="Mechanical behaviour of panels S2\u2013S4: "
                 "(a) stress\u2013strain curves (representative replicate; "
                 "failure at \u03b5\u1da0 \u2248 16%) and "
                 "(b) mean unconfined compressive strength "
                 "(error bars = \u00b1 1 SD; S1 excluded).")

heading(doc, "3.4.  Thermal Conductivity", level=2)
P(doc,
  "Thermal conductivity results are presented in Table 4 and Figures 9\u201310. "
  "Both the fine binder-less panel S2 (k = 0.0577 W m\u207b\u00b9 K\u207b\u00b9) "
  "and the 90:10 panel S3 (k = 0.0608 W m\u207b\u00b9 K\u207b\u00b9) fall below "
  "the ASTM C168 insulation-grade threshold of 0.065 W m\u207b\u00b9 K\u207b\u00b9. "
  "All S\u1d67\u1d7a values were below 2.0, confirming acceptable measurement "
  "quality. The S2 and S3 conductivities are competitive with the "
  "0.047\u20130.065 W m\u207b\u00b9 K\u207b\u00b9 range of binder-less WH petiole "
  "boards (Salas-Ruiz et al., 2019), and are within the same insulation-grade "
  "band as cellulose-fibre insulation (\u224840\u201350 mW m\u207b\u00b9 K\u207b\u00b9) "
  "and hemp boards (Kym\u00e4l\u00e4inen and Sj\u00f6berg, 2008).")
P(doc,
  "The high-starch panel S4 exhibited a three-fold higher conductivity "
  "(0.1846 W m\u207b\u00b9 K\u207b\u00b9), placing it outside the insulation grade. "
  "Two mechanisms explain this: (i) the dense continuous starch matrix fills "
  "air voids that are primarily responsible for insulation; and (ii) S4's very "
  "high moisture content (41.18%) means a large fraction of pore space is "
  "occupied by water rather than air. Since water conducts heat roughly 23 times "
  "more effectively than air, even a modest increase in moisture dramatically "
  "raises the effective conductivity \u2014 the same moisture-driven degradation "
  "reported by Jeon et al. (2017) for glass wool and by Zhou et al. (2022) for "
  "straw boards.")

tbl_caption(doc, "Table 4",
            "Thermal conductivity of the bio-composite panels "
            "(KD2 Pro TR-3, ASTM D5334; mean \u00b1 SD, n = 3). "
            "S1 was not tested (fractured on demoulding).")
make_table(doc,
    ["Sample", "Composition",
     "k (W m\u207b\u00b9 K\u207b\u00b9)", "S\u1d67\u1d7a",
     "Classification (ASTM C168)"],
    [["S2", "Fine, no binder",
      "0.0577 \u00b1 0.0012", "0.44",
      "Insulation grade (k < 0.065)"],
     ["S3", "Fine, 90:10 starch",
      "0.0608 \u00b1 0.0015", "0.32",
      "Insulation grade (k < 0.065)"],
     ["S4", "Fine, 70:30 starch",
      "0.1846 \u00b1 0.0041", "1.21",
      "Non-insulation grade"]],
    col_widths=[0.60, 1.48, 1.65, 0.65, 2.27], font_size=10.5)

# Thermal test photographs
insert_image_row(doc,
    [(f'{PHOTOS}/thermal_test_nobinder.jpg', '(a) S2 \u2013 binder-less'),
     (f'{PHOTOS}/thermal_test_9010.jpg',     '(b) S3 \u2013 90:10 starch'),
     (f'{PHOTOS}/thermal_test_7030.jpg',     '(c) S4 \u2013 70:30 starch')],
    total_width_in=6.0,
    label="Figure 9",
    caption_text="KD2 Pro TR-3 probe embedded in the bio-composite panels "
                 "during thermal conductivity measurement: (a) S2 (binder-less), "
                 "(b) S3 (90:10 starch) and (c) S4 (70:30 starch).")

insert_image(doc, f'{CHARTS}/Fig18_Thermal_Conductivity.png', 4.6,
    "Figure 10",
    "Mean thermal conductivity of the bio-composite panels. The dashed line "
    "marks the ASTM C168 insulation-grade threshold "
    "(k = 0.065 W m\u207b\u00b9 K\u207b\u00b9); error bars = \u00b1 1 SD; "
    "S1 excluded.")


heading(doc, "3.5.  Microstructural Characterisation (SEM)", level=2)
P(doc,
  "SEM micrographs of the fine binder-less panel S2 (Figure 11) and the 90:10 "
  "corn-starch panel S3 (Figure 12) provide direct microstructural evidence "
  "correlating the macroscopic performance differences to the underlying "
  "fibre-matrix architecture.")
P(doc,
  "In the binder-less panel S2, the micrographs reveal a loosely packed network "
  "of irregular cellulosic fibre bundles with rough, corrugated surfaces and no "
  "bridging adhesive matrix. Inter-fibre contacts are maintained solely by "
  "physical interlocking and surface friction. Large open macro-pores are "
  "distributed throughout the cross-section. At higher magnification, the "
  "characteristic vascular bundle architecture of the water-hyacinth petiole "
  "\u2014 large aerenchyma air channels surrounded by thin fibre walls \u2014 is "
  "clearly visible. This highly porous, air-rich microstructure is the direct "
  "physical origin of S2's low thermal conductivity "
  "(0.0577 W m\u207b\u00b9 K\u207b\u00b9). The absence of a cohesive adhesive "
  "matrix explains the moderate compressive strength (94.4 kPa): under axial "
  "load, fibres slip at contact points rather than fracturing through a "
  "continuous matrix.")

# S2 SEM – 5 images in 2+2+1 layout
insert_image_row(doc,
    [(f'{PHOTOS}/sem_s1_00.jpg', '(a)'),
     (f'{PHOTOS}/sem_s1_01.jpg', '(b)')], total_width_in=5.6)
insert_image_row(doc,
    [(f'{PHOTOS}/sem_s1_02.jpg', '(c)'),
     (f'{PHOTOS}/sem_s1_03.jpg', '(d)')], total_width_in=5.6)
insert_image(doc, f'{PHOTOS}/sem_s1_04.jpg', 2.8,
    "Figure 11",
    "SEM micrographs of the fine binder-less panel S2 at increasing "
    "magnification: (a)\u2013(b) low magnification showing overall fibre "
    "arrangement and large open macro-pores; (c)\u2013(d) intermediate "
    "magnification showing individual fibre bundle morphology and open "
    "inter-fibre void space; (e) high magnification revealing the vascular "
    "bundle architecture and corrugated fibre surface with no binder film.")

P(doc,
  "The 90:10 starch-bound panel S3 shows a markedly different microstructure. "
  "A continuous gelatinised corn-starch matrix is clearly visible coating the "
  "fibre surfaces and forming bridges at inter-fibre contact points. The large "
  "open macro-pores seen in S2 are substantially reduced in size and number, "
  "consistent with S3's higher bulk density (0.992 vs 0.900 g cm\u207b\u00b3). "
  "At high magnification the starch\u2013fibre interface appears well-bonded "
  "with no visible delamination, indicating good adhesion between the "
  "gelatinised starch film and the cellulosic fibre surface. This coherent "
  "load-bearing matrix explains the 97% increase in compressive strength of S3 "
  "relative to S2. The partial retention of open pore space is also consistent "
  "with S3 remaining insulation-grade: 10 wt% starch displaces some insulating "
  "air, but not enough to raise k above the 0.065 W m\u207b\u00b9 K\u207b\u00b9 "
  "threshold. The SEM evidence therefore corroborates the optimality of the "
  "90:10 formulation: starch builds a continuous load-bearing matrix without "
  "compromising thermal performance.")

# S3 SEM – 6 images in 2+2+2 layout
insert_image_row(doc,
    [(f'{PHOTOS}/sem_s2_05.jpg', '(a)'),
     (f'{PHOTOS}/sem_s2_06.jpg', '(b)')], total_width_in=5.6)
insert_image_row(doc,
    [(f'{PHOTOS}/sem_s2_07.jpg', '(c)'),
     (f'{PHOTOS}/sem_s2_08.jpg', '(d)')], total_width_in=5.6)
insert_image_row(doc,
    [(f'{PHOTOS}/sem_s2_09.jpg', '(e)'),
     (f'{PHOTOS}/sem_s2_10.jpg', '(f)')],
    total_width_in=5.6,
    label="Figure 12",
    caption_text="SEM micrographs of the 90:10 corn-starch panel S3 at "
                 "increasing magnification: (a)\u2013(b) low magnification "
                 "showing reduced macro-porosity and starch-coated fibre bundles; "
                 "(c)\u2013(d) intermediate magnification showing the continuous "
                 "starch matrix bridging fibres at contact points; (e)\u2013(f) "
                 "high magnification revealing the starch\u2013fibre interface "
                 "morphology with good adhesion and no delamination.")


heading(doc, "3.6.  Optimal Formulation and Comparison with Reported Materials",
        level=2)
P(doc,
  "Considering the combined requirements of insulation-grade thermal conductivity, "
  "adequate compressive strength and a well-bonded microstructure, the 90:10 "
  "biomass:starch panel S3 emerges as the best-balanced formulation. It delivers "
  "the highest compressive strength (186.0 kPa) while remaining insulation-grade "
  "(k = 0.0608 W m\u207b\u00b9 K\u207b\u00b9), and SEM confirms a well-bonded "
  "continuous starch matrix. The binder-less panel S2 is also insulation-grade "
  "but offers only half the compressive strength. The high-starch panel S4 offers "
  "the lowest water absorption (280%) but fails on both thermal and mechanical "
  "criteria. The coarse panel S1 fractured on demoulding, demonstrating that "
  "fine particle size is non-negotiable for panel integrity.")
P(doc,
  "Table 5 benchmarks the present panels against conventional insulants and "
  "reported bio-composites. The conductivities of S2 and S3 are higher than "
  "glass wool (\u224834 mW m\u207b\u00b9 K\u207b\u00b9) and EPS "
  "(35\u201340 mW m\u207b\u00b9 K\u207b\u00b9), but fall within the same "
  "insulation-grade band and are competitive with binder-less WH boards "
  "(Salas-Ruiz et al., 2019), hemp boards (Kym\u00e4l\u00e4inen and Sj\u00f6berg, "
  "2008), corn-cob boards (Pinto et al., 2021) and rice-straw panels "
  "(Zhou et al., 2022). Crucially, S3 achieves this performance using only a "
  "biodegradable corn-starch binder derived from an invasive weed waste stream, "
  "with no synthetic resin, cement or petrochemical additive \u2014 a clear "
  "embodied-carbon and end-of-life advantage over epoxy-based "
  "(Anjani et al., 2023) and cement-based (Philip and Rakendu, 2022) alternatives.")

tbl_caption(doc, "Table 5",
            "Comparison of the present bio-composite panels with conventional "
            "insulants and reported water-hyacinth and other bio-based composites.")
make_table(doc,
    ["Material",
     "k (W m\u207b\u00b9 K\u207b\u00b9)",
     "Binder / matrix",
     "Reference"],
    [["Glass wool",
      "\u2248 0.034",
      "Inorganic (mineral)",
      "Jeon et al. (2017)"],
     ["EPS foam",
      "0.035\u20130.040",
      "Polystyrene (synthetic)",
      "Typical value"],
     ["Hemp / flax board",
      "0.040\u20130.060",
      "None / mineral binder",
      "Kym\u00e4l\u00e4inen & Sj\u00f6berg (2008)"],
     ["Corn-cob board",
      "0.046\u20130.052",
      "None",
      "Pinto et al. (2021)"],
     ["Rice-straw board",
      "0.050\u20130.070",
      "Natural binder",
      "Zhou et al. (2022)"],
     ["WH binder-less board",
      "0.047\u20130.065",
      "None",
      "Salas-Ruiz et al. (2019)"],
     ["WH\u2013rubber latex board",
      "\u2248 0.055",
      "Natural rubber latex",
      "Jaktorn & Jiajitsawat (2021)"],
     ["WH\u2013cement board",
      "0.0765",
      "Portland cement",
      "Philip & Rakendu (2022)"],
     ["WH\u2013bagasse\u2013epoxy",
      "0.1987",
      "Epoxy resin (synthetic)",
      "Anjani et al. (2023)"],
     ["S2 (this study)",
      "0.0577",
      "None (water only)",
      "\u2014"],
     ["S3 (this study)",
      "0.0608",
      "Corn starch 10 wt%",
      "\u2014"],
     ["S4 (this study)",
      "0.1846",
      "Corn starch 30 wt%",
      "\u2014"]],
    col_widths=[1.90, 1.32, 1.60, 1.83], font_size=10.0)

print("Section 3 done.")


# ══════════════════════════════ 4. CONCLUSIONS ══════════════════════════════

heading(doc, "4.  Conclusions")
P(doc,
  "This study demonstrated the technical feasibility of valorising invasive "
  "aquatic weeds from Dal Lake, Srinagar, into fully bio-based insulation panels "
  "using food-grade corn starch as the sole binder. The following principal "
  "conclusions are drawn:")

CONCLUSIONS = [
    "Particle size is decisive. The coarse binder-less panel (S1) fractured on "
    "demoulding due to insufficient inter-particle cohesion, whereas all "
    "fine-particle panels (1.0\u20131.5 mm) formed intact specimens, confirming "
    "that fine particle size is a necessary condition for structural integrity.",

    "An optimal binder content of 10 wt% corn starch was identified. The 90:10 "
    "panel S3 achieved the highest unconfined compressive strength (186.0 kPa), "
    "a 97% improvement over the binder-less fine panel S2 (94.4 kPa). SEM "
    "confirmed a continuous, well-adhered gelatinised starch matrix bridging "
    "fibres at inter-particle contact points in S3.",

    "Two formulations qualified as insulation-grade per ASTM C168: S2 "
    "(k = 0.0577 W m\u207b\u00b9 K\u207b\u00b9) and S3 "
    "(k = 0.0608 W m\u207b\u00b9 K\u207b\u00b9). Their low conductivities "
    "originate from the high volume fraction of air in the porous fibre network, "
    "as directly confirmed by SEM.",

    "Excess starch is detrimental. The 70:30 panel S4 showed the lowest "
    "compressive strength (26.8 kPa) and a three-fold higher conductivity "
    "(0.1846 W m\u207b\u00b9 K\u207b\u00b9), attributable to void filling by "
    "the dense starch matrix and a high residual moisture content (41.18%) "
    "that displaces insulating air with conductive water.",

    "Water absorption was high for all panels (280\u2013658%) but decreased "
    "with increasing binder fraction. Moisture resistance is the principal "
    "practical limitation and must be addressed by hydrophobic surface treatment "
    "before field deployment.",

    "The 90:10 biomass:starch panel S3 is the best-balanced formulation, "
    "uniquely combining insulation-grade thermal conductivity with the highest "
    "mechanical strength and a well-bonded microstructure. It represents a "
    "viable, biodegradable and fully bio-based alternative to EPS and PU foam, "
    "derived from a freely available invasive-weed waste stream that must already "
    "be removed to protect Dal Lake.",
]

for bullet in CONCLUSIONS:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(0)
    r = bp.add_run(bullet)
    r.font.name = FONT; r.font.size = Pt(12)

P(doc,
  "Future work should address: (i) hydrophobic surface treatment using beeswax, "
  "silicone or linseed-oil coatings; (ii) mild alkali fibre pre-treatment to "
  "improve surface adhesion and moisture resistance; (iii) eco-friendly "
  "flame-retardant additives; (iv) hot-pressing trials to increase density; "
  "(v) FTIR characterisation of the starch\u2013fibre chemical interface; and "
  "(vi) a comparative life-cycle assessment against EPS, PU foam and glass wool "
  "to quantify the embodied-carbon advantage of this fully bio-based system.",
  space_before=10)

print("Section 4 done.")


# ════════════════════════════════ DECLARATIONS ═══════════════════════════════

heading(doc, "CRediT Authorship Contribution Statement", level=2,
        space_before=18, space_after=4)
P(doc,
  "Suchetan Karloopia: Conceptualization, Methodology, Investigation, Writing "
  "\u2013 original draft, Visualization. "
  "Miran Haider: Investigation, Validation, Data curation. "
  "Akshita Sen: Investigation, Formal analysis, Visualization. "
  "Fasil Qayoom Mir: Supervision, Conceptualization, Writing \u2013 review & "
  "editing, Resources, Project administration.",
  size=12)

heading(doc, "Declaration of Competing Interest", level=2,
        space_before=14, space_after=4)
P(doc,
  "The authors declare that they have no known competing financial interests or "
  "personal relationships that could have appeared to influence the work reported "
  "in this paper.",
  size=12)

heading(doc, "Data Availability", level=2, space_before=14, space_after=4)
P(doc,
  "The raw experimental data supporting the conclusions of this article will be "
  "made available by the corresponding author on reasonable request.",
  size=12)

heading(doc, "Acknowledgements", level=2, space_before=14, space_after=4)
P(doc,
  "The authors gratefully acknowledge the Department of Chemical Engineering, "
  "NIT Srinagar, for laboratory facilities; the Department of Chemistry, NIT "
  "Srinagar, for access to the mixer-grinder; the Department of Civil "
  "Engineering, NIT Srinagar, for the Baker Type K12 unconfined compression "
  "apparatus; and the technical staff who facilitated the KD2 Pro "
  "thermal-conductivity measurements. The authors also thank the Jammu & Kashmir "
  "Lakes and Waterways Development Authority for permission to collect aquatic "
  "biomass samples from Dal Lake.",
  size=12)

print("Declarations done.")


# ══════════════════════════════════ REFERENCES ═══════════════════════════════

heading(doc, "References", level=1, space_before=20, space_after=6)

REFS = [
    "Abral, H., Kadriadi, D., Rodianus, A., Mistar, P., Sapuan, S.M., "
    "Zainudin, E.S., 2014. Mechanical properties of water hyacinth "
    "fibers\u2013polyester composites before and after immersion in water. "
    "Mater. Des. 58, 125\u2013129. "
    "https://doi.org/10.1016/j.matdes.2014.01.043",

    "Anjani, A., Iskandar, B., Aziz, M., 2023. The utilization of composite "
    "material: water hyacinth and sugarcane bagasse fiber\u2013epoxy for cool "
    "box thermal insulation. J. Energy Mech. Mater. Manuf. Eng. 8 (1), "
    "29\u201338. https://doi.org/10.22219/jemmme.v8i1.25010",

    "Asdrubali, F., D\u2019Alessandro, F., Schiavoni, S., 2015. A review of "
    "unconventional sustainable building insulation materials. Sustain. Mater. "
    "Technol. 4, 1\u201317. https://doi.org/10.1016/j.susmat.2015.05.002",

    "Chaireh, P., Meethong, N., Khomwaen, P., 2020. Novel composite foam made "
    "from starch and water hyacinth with beeswax coating for food packaging "
    "applications. Int. J. Biol. Macromol. 165, 1382\u20131391. "
    "https://doi.org/10.1016/j.ijbiomac.2020.09.243",

    "Cosentino, L., Fernandes, P., Mateus, R., 2023. A review of natural "
    "bio-based insulation materials. Energies 16 (13), 4926. "
    "https://doi.org/10.3390/en16134926",

    "Jaktorn, C., Jiajitsawat, S., 2021. Production of thermal insulator from "
    "water hyacinth fiber and natural rubber latex. J. Ecol. Eng. 22 (7), "
    "134\u2013141. https://doi.org/10.12911/22998993/138736",

    "Jeon, J., Park, S., Kim, S., 2017. A study on insulation characteristics "
    "of glass wool and mineral wool coated with a polysiloxane agent. Adv. "
    "Mater. Sci. Eng. 2017, 3938965. https://doi.org/10.1155/2017/3938965",

    "Kym\u00e4l\u00e4inen, H.-R., Sj\u00f6berg, A.-M., 2008. Flax and hemp "
    "fibres as raw materials for thermal insulation. Build. Environ. 43 (7), "
    "1261\u20131269. https://doi.org/10.1016/j.buildenv.2007.03.006",

    "L\u00f3pez, O.V., Zaritzky, N.E., Grossmann, M.V.E., Garc\u00eda, M.A., "
    "2013. Acetylated and native corn starch blend films produced by casting. "
    "J. Food Eng. 116 (2), 286\u2013297. "
    "https://doi.org/10.1016/j.jfoodeng.2012.12.002",

    "Madhu, P., Praveenkumara, J., Mavinkere Rangappa, S., Siengchin, S., 2019. "
    "A comprehensive review on synthesis and characterization of agro-industrial "
    "waste-based bio-composites for construction and structural applications. "
    "J. Clean. Prod. 246, 119003. "
    "https://doi.org/10.1016/j.jclepro.2019.119003",

    "Mitra, S., Bhatt, D., Bhatt, S., 2020. Assessment of Dal Lake, Kashmir: "
    "a review on water quality, biodiversity, threats and restoration measures. "
    "Environ. Monit. Assess. 192, 774. "
    "https://doi.org/10.1007/s10661-020-08723-8",

    "Oushabi, A., Sair, S., Abboud, Y., Tanane, O., El Bouari, A., 2022. "
    "Thermal and mechanical characterization of alkali-treated sugarcane "
    "bagasse-reinforced thermoset composites. South Afr. J. Chem. Eng. 40, "
    "104\u2013112. https://doi.org/10.1016/j.sajce.2022.02.006",

    "Pawlowski, K., Strzalkowska, A., Chojnacka, B., 2025. Exploring "
    "advancements in bio-based composites for thermal insulation: a systematic "
    "review. Sustainability 17 (3), 1143. "
    "https://doi.org/10.3390/su17031143",

    "Philip, S., Rakendu, R., 2022. Thermal insulation materials based on "
    "water hyacinth for application in sustainable buildings. Mater. Today "
    "Proc. 57, 1863\u20131867. https://doi.org/10.1016/j.matpr.2022.01.062",

    "Pinto, J., Pereira, E., Tavares, A., Ferreira, V.M., 2021. Corn\u2019s "
    "cob as a potential ecological thermal insulation material. Constr. Build. "
    "Mater. 277, 122282. "
    "https://doi.org/10.1016/j.conbuildmat.2021.122282",

    "Ruiz-G\u00f3mez, N.A., Fonseca-Florido, H.A., Rios-Soberanis, C.R., "
    "Arag\u00f3n-Pi\u00f1a, A., Castillo-Atoche, A.C., 2021. Cassava "
    "starch-based films reinforced with sisal fibres: moisture and UV "
    "radiation effects on degradation. Polymers 13 (4), 600. "
    "https://doi.org/10.3390/polym13040600",

    "Salas-Ruiz, A., Barbero-Barrera, M. del M., Ruiz-T\u00e9llez, T., 2019. "
    "Microstructural and thermo-physical characterization of a water hyacinth "
    "petiole for thermal insulation particle board manufacture. Materials 12 "
    "(4), 560. https://doi.org/10.3390/ma12040560",

    "Tanobe, V.O.A., Sydenstricker, T.H.D., Munaro, M., Amico, S.C., 2005. "
    "A comprehensive characterization of chemically treated Brazilian "
    "sponge-gourds (Luffa cylindrica). Polym. Test. 24 (4), 474\u2013482. "
    "https://doi.org/10.1016/j.polymertesting.2004.12.004",

    "Tazmeen, T., Mir, F.Q., 2024. Sustainability through materials: a review "
    "of green options in construction. Results Surf. Interfaces 14, 100206. "
    "https://doi.org/10.1016/j.rsurfi.2024.100206",

    "Wang, X., Li, Z., Shi, H., Yu, Y., 2023. Natural pineapple-leaf fibre: "
    "a promising material for high-performance composites. Ind. Crops Prod. "
    "195, 116447. https://doi.org/10.1016/j.indcrop.2023.116447",

    "Zhou, Y., Trabelsi, A., El Mankibi, M., 2022. Hygrothermal properties "
    "of insulation materials from rice straw and natural binders for "
    "buildings. Polymers 14 (9), 1735. "
    "https://doi.org/10.3390/polym14091735",
]

for ref in REFS:
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    rp.paragraph_format.space_before      = Pt(0)
    rp.paragraph_format.space_after       = Pt(0)
    rp.paragraph_format.left_indent       = Inches(0.4)
    rp.paragraph_format.first_line_indent = Inches(-0.4)
    r = rp.add_run(ref)
    r.font.name = FONT; r.font.size = Pt(12)

print("References done.")

# ─────────────────────────────────── save ───────────────────────────────────
doc.save(OUTPUT)
print("SAVED:", OUTPUT)
